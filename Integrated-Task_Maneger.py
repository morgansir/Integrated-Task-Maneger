
import os
import sys
import math
import pickle
import logging
import sqlite3
import getpass
import hashlib
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import psutil
import ctypes
import csv
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import numpy as np
from sklearn.ensemble import RandomForestClassifier
import yaml  # requires PyYAML (install via pip install pyyaml)

# إعداد ملف تسجيل الأخطاء
logging.basicConfig(
    filename='integrated_app_errors.log',
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# إعداد ألوان الواجهة -- ألوان متدرجة وعصرية
COLORS = {
    'safe': '#d4edda',
    'warning': '#fff3cd',
    'danger': '#f8d7da',
    'header': '#2c3e50',
    'background': '#e6f7ff',  # خلفية متدرجة بدلاً من الأبيض
    'stats_bg': '#2c3e50',
    'stats_fg': 'white',
    'button_alert': '#dc3545',
    'button_normal': '#007bff',
    'gradient1': '#2c3e50',
    'gradient2': '#c2e9fb',
    'gradient3': '#a1c4fd',
    'high_usage': '#FF6347'
}

# بيانات Threat Intelligence الافتراضية
DEFAULT_THREAT_INTEL = {
    "suspicious_names": ["svchost.exe", "winlogon.exe", "rundll32.exe", "mshta.exe", "wscript.exe"],
    "suspicious_keywords": ["malware", "trojan", "keylogger"]
}

# قواعد SIGMA الافتراضية (للحالة الأولى في حال عدم وجود بيانات في قاعدة البيانات)

DEFAULT_SIGMA_RULES = {
    "ransomware": {
        "description": "قاعدة تكشف سلوك برامج الفدية",
        "mitre": "T1486"  # Data Encrypted for Impact
    },
    "spyware": {
        "description": "قاعدة تكشف أنشطة برامج التجسس",
        "mitre": "T1127"  # Shared Modules
    },
    "keylogger": {
        "description": "قاعدة تكشف استخدام أدوات تسجيل ضغطات المفاتيح",
        "mitre": "T1056.001"  # Keylogging
    },
    "trojan": {
        "description": "قاعدة تكشف سلوك التروجان",
        "mitre": "T1219"  # Remote Access Tools
    }
}

HIGH_RESOURCE_THRESHOLD = {
    'cpu': 80.0,
    'memory': 70.0
}

ANOMALY_THRESHOLD = 2.0

MODEL_FILENAME = "process_model.pkl"
DB_FILENAME = "system_processes.db"

def is_admin():
    """يتحقق مما إذا كان البرنامج يعمل بصلاحيات المدير على ويندوز أو لينكس"""
    if sys.platform == "win32":
        try:
            return ctypes.windll.shell32.IsUserAnAdmin()
        except Exception:
            return False
    else:
        return os.getuid() == 0

def load_checkbox_images(root):
    """تحميل صور مربعات الاختيار باستخدام بيانات Base64"""
    checked_data = '''
        R0lGODlhEAAQAPIAAP///wAAAMLCwkJCQgAAAAAAACH5BAAAAAAALAAAAAAQABAAAAM6SLrc/jDKSesyymKrG6lt6D/5ooSFQA7
    '''
    unchecked_data = '''
        R0lGODlhEAAQAPIAAP///wAAAMLCwkJCQgAAAAAAACH5BAAAAAAALAAAAAAQABAAAAM+SLrc/jDKSesyymKrG6lt6D/5ooSFQA7
    '''
    checkbox_checked = tk.PhotoImage(data=checked_data)
    checkbox_unchecked = tk.PhotoImage(data=unchecked_data)

    return checkbox_checked, checkbox_unchecked

def load_threat_intel(file_path=None):
    """إرجاع بيانات Threat Intelligence المُدمجة داخل الكود"""
    return DEFAULT_THREAT_INTEL

def initialize_database():
    """تهيئة قاعدة البيانات وإنشاء الجداول إذا لم تكن موجودة"""
    conn = sqlite3.connect(DB_FILENAME)
    c = conn.cursor()
    # جدول الكلمات المفتاحية
    c.execute('''
        CREATE TABLE IF NOT EXISTS keywords (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            keyword TEXT UNIQUE
        )
    ''')
    # جدول تاريخ العمليات
    c.execute('''
        CREATE TABLE IF NOT EXISTS process_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp DATETIME,
            process_name TEXT,
            pid INTEGER,
            cpu_percent REAL,
            memory_percent REAL,
            path TEXT,
            command_line TEXT,
            status TEXT,
            suspicion_score REAL,
            reason TEXT,
            user_info TEXT
        )
    ''')
    # جدول قواعد SIGMA
    c.execute('''
        CREATE TABLE IF NOT EXISTS sigma_rules (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rule_key TEXT UNIQUE,
            description TEXT,
            mitre TEXT
        )
    ''')
    # إدخال القواعد الافتراضية في حال كانت فارغة
    c.execute("SELECT COUNT(*) FROM sigma_rules")
    count = c.fetchone()[0]
    if count == 0:
        for key, data in DEFAULT_SIGMA_RULES.items():
            c.execute("INSERT OR IGNORE INTO sigma_rules (rule_key, description, mitre) VALUES (?, ?, ?)",
                      (key, data['description'], data['mitre']))
    conn.commit()

    conn.close()

class IntegratedTaskManager:
    def __init__(self, root):
        self.root = root
        self.root.title("Integrated Task Manager")
        self.configure_window_size()
        self.create_header()

        self.checkbox_checked, self.checkbox_unchecked = load_checkbox_images(self.root)

        if not is_admin():
            messagebox.showwarning("Permission Warning", "هذا البرنامج يجب تشغيله بصلاحيات المدير!")
            sys.exit(1)

        try:
            self.conn = sqlite3.connect(DB_FILENAME, check_same_thread=False)
            self.c = self.conn.cursor()
            initialize_database()
            self.c.execute('CREATE TABLE IF NOT EXISTS keywords (id INTEGER PRIMARY KEY AUTOINCREMENT, keyword TEXT UNIQUE)')
            self.conn.commit()
        except Exception as e:
            self.show_error("Database Error", e)
            sys.exit(1)

        self.threat_intel = load_threat_intel()
        # بدءاً من القائمة الافتراضية
        self.suspicious_keywords = self.threat_intel.get("suspicious_keywords", [])
        self.load_keywords()
        if "keylogger" not in self.suspicious_keywords:
            self.suspicious_keywords.append("keylogger")

        # تحميل نموذج التعلم الآلي
        self.load_ml_model()

        # تحميل قواعد SIGMA من قاعدة البيانات
        self.load_sigma_rules()

        self.setup_styles()
        self.create_widgets()

        self.root.after(10000, self.safe_update)

    def load_keywords(self):
        """تحميل الكلمات المخزنة من قاعدة البيانات وإضافتها إلى القائمة"""
        try:
            self.c.execute("SELECT keyword FROM keywords")
            rows = self.c.fetchall()
            for row in rows:
                keyword = row[0]

                if keyword not in self.suspicious_keywords:
                    self.suspicious_keywords.append(keyword)
        except Exception as e:
            logging.error(f"Error loading keywords: {str(e)}")

    def load_sigma_rules(self):
        """تحميل قواعد SIGMA من قاعدة البيانات إلى قاموس محلي"""
        self.sigma_rules = {}
        try:
            self.c.execute("SELECT rule_key, description, mitre FROM sigma_rules")
            rows = self.c.fetchall()
            for row in rows:
                key, desc, mitre = row
                self.sigma_rules[key] = {"description": desc, "mitre": mitre}
        except Exception as e:
            logging.error(f"Error loading sigma rules: {str(e)}")

    def load_ml_model(self):
        """تحميل نموذج التعلم الآلي، أو تدريبه إذا لم يكن موجوداً"""
        if os.path.exists(MODEL_FILENAME):
            try:
                with open(MODEL_FILENAME, "rb") as f:
                    self.ml_model = pickle.load(f)
            except Exception as e:
                logging.error(f"Error loading ML model: {str(e)}")
                self.train_dummy_model()
        else:
            self.train_dummy_model()

    def train_dummy_model(self):
        """تدريب نموذج بسيط باستخدام بيانات عشوائية للتوضيح"""
        try:
            X_train = np.random.rand(100, 2) * 100  # قيم عشوائية بين 0 و100
            y_train = np.random.randint(0, 2, 100)    # 0: طبيعي، 1: مشبوه
            model = RandomForestClassifier(n_estimators=100, random_state=42)
            model.fit(X_train, y_train)
            self.ml_model = model
            with open(MODEL_FILENAME, "wb") as f:
                pickle.dump(model, f)
        except Exception as e:
            logging.error(f"Error training ML model: {str(e)}")
            self.ml_model = None

    def extract_features(self, proc_info):
        """استخراج ميزات عملية بسيطة (cpu_percent و memory_percent) لاستخدامها مع نموذج التعلم الآلي"""
        cpu = proc_info.get('cpu_percent', 0)
        mem = proc_info.get('memory_percent', 0)
        return np.array([[cpu, mem]])

    def update_dynamic_threat_intel(self):
        """محاكاة تحديث Threat Intelligence ديناميكيًا"""

        new_keywords = ["ransomware", "spyware"]
        for kw in new_keywords:
            if kw not in self.threat_intel["suspicious_keywords"]:
                self.threat_intel["suspicious_keywords"].append(kw)
                try:
                    self.c.execute("INSERT OR IGNORE INTO keywords (keyword) VALUES (?)", (kw,))
                    self.conn.commit()
                    if kw not in self.suspicious_keywords:
                        self.suspicious_keywords.append(kw)
                except Exception as e:
                    logging.error(f"Error updating dynamic threat intel: {str(e)}")

    def configure_window_size(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = int(screen_width * 0.9)
        window_height = int(screen_height * 0.85)
        x_pos = (screen_width - window_width) // 2
        y_pos = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x_pos}+{y_pos}")
        self.root.minsize(1200, 700)

    def create_header(self):
        """إنشاء رأس بتدرج ثلاثي الأبعاد لعنوان التطبيق"""
        header_frame = tk.Frame(self.root, bg=COLORS['gradient1'])
        header_frame.pack(fill=tk.X)
        canvas = tk.Canvas(header_frame, height=60, bg=COLORS['gradient1'], highlightthickness=0)
        canvas.pack(fill=tk.X)
        (r1, g1, b1) = self.root.winfo_rgb(COLORS['gradient1'])
        (r2, g2, b2) = self.root.winfo_rgb(COLORS['gradient3'])
        for i in range(60):
            nr = int(r1 + (r2 - r1) * i / 60)
            ng = int(g1 + (g2 - g1) * i / 60)
            nb = int(b1 + (b2 - b1) * i / 60)
            color = f'#{nr//256:02x}{ng//256:02x}{nb//256:02x}'
            canvas.create_line(0, i, self.root.winfo_screenwidth(), i, fill=color)
        canvas.create_text(self.root.winfo_screenwidth()//2, 30, text="Integrated Task Manager", fill="white", font=("Segoe UI", 20, "bold"))

    def setup_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')
        self.style.configure("TFrame", background=COLORS['background'])
        self.style.configure("TLabel", background=COLORS['background'], foreground=COLORS['header'], font=("Segoe UI", 10))
        self.style.configure('Stats.TLabelframe',
                             background=COLORS['stats_bg'],
                             foreground=COLORS['stats_fg'],
                             borderwidth=2,
                             relief='groove')
        self.style.configure('StatsLabel.TLabel',

                             font=("Segoe UI", 11, "bold"),
                             foreground='white',
                             background=COLORS['stats_bg'],
                             padding=5)
        self.style.configure('Normal.TButton',
                             font=("Segoe UI", 10),
                             foreground='white',
                             background=COLORS['button_normal'],
                             padding=8,
                             relief="raised",
                             borderwidth=2)
        self.style.map('Normal.TButton', background=[("active", "#0056b3")])
        self.style.configure('Alert.TButton',
                             font=("Segoe UI", 10, "bold"),
                             foreground='white',
                             background=COLORS['button_alert'],
                             padding=8,
                             relief="raised",
                             borderwidth=2)
        self.style.map('Alert.TButton', background=[("active", "#a71d2a")])
        self.style.configure("Treeview", background="#ffffff", fieldbackground="#ffffff", foreground="black")
        self.style.configure("Treeview.Heading", background=COLORS['gradient1'], foreground="white", font=("Segoe UI", 10, "bold"))
        self.style.configure("TNotebook", background=COLORS['background'])
        self.style.configure("TNotebook.Tab", background=COLORS['gradient2'], foreground="white", padding=[10, 5], font=("Segoe UI", 10, "bold"))
        self.style.map("TNotebook.Tab", background=[("selected", COLORS['gradient3'])])

    def create_widgets(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        self.comprehensive_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.comprehensive_frame, text="Comprehensive Scan")
        self.create_comprehensive_tab(self.comprehensive_frame)

        self.monitor_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.monitor_frame, text="Process Monitor")
        self.create_monitor_tab(self.monitor_frame)

        self.keyword_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.keyword_frame, text="Keyword Scan")
        self.create_keyword_tab(self.keyword_frame)

        self.history_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.history_frame, text="History")
        self.create_history_tab(self.history_frame)

        self.anomaly_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.anomaly_frame, text="Anomaly Detection")
        self.create_anomaly_tab(self.anomaly_frame)


        # New Tab for Sigma Rules Updates
        self.sigma_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.sigma_frame, text="Sigma Rules")
        self.create_sigma_tab(self.sigma_frame)

    def create_sigma_tab(self, parent):
        """واجهة لتحديث وإدارة قواعد SIGMA"""
        top_frame = ttk.Frame(parent)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(top_frame, text="قواعد SIGMA الحالية:").pack(side=tk.LEFT, padx=5)
        refresh_btn = ttk.Button(top_frame, text="تحديث العرض", command=self.refresh_sigma_tree, style='Normal.TButton')
        refresh_btn.pack(side=tk.LEFT, padx=5)
        import_btn = ttk.Button(top_frame, text="استيراد من ملف", command=self.import_sigma_rules, style='Normal.TButton')
        import_btn.pack(side=tk.RIGHT, padx=5)

        tree_frame = ttk.Frame(parent)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        columns = ("rule_key", "description", "mitre")
        self.sigma_tree = ttk.Treeview(tree_frame, columns=columns, show='headings')
        for col, header in zip(columns, ["المفتاح", "الوصف", "MITRE"]):
            self.sigma_tree.heading(col, text=header)
            self.sigma_tree.column(col, anchor=tk.CENTER, width=150)
        self.sigma_tree.pack(fill=tk.BOTH, expand=True)

        form_frame = ttk.Frame(parent)
        form_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(form_frame, text="المفتاح:").grid(row=0, column=0, padx=5, pady=2, sticky=tk.E)
        self.rule_key_entry = ttk.Entry(form_frame)
        self.rule_key_entry.grid(row=0, column=1, padx=5, pady=2, sticky=tk.W)
        ttk.Label(form_frame, text="الوصف:").grid(row=1, column=0, padx=5, pady=2, sticky=tk.E)
        self.rule_desc_entry = ttk.Entry(form_frame)
        self.rule_desc_entry.grid(row=1, column=1, padx=5, pady=2, sticky=tk.W)
        ttk.Label(form_frame, text="MITRE:").grid(row=2, column=0, padx=5, pady=2, sticky=tk.E)
        self.rule_mitre_entry = ttk.Entry(form_frame)
        self.rule_mitre_entry.grid(row=2, column=1, padx=5, pady=2, sticky=tk.W)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        add_rule_btn = ttk.Button(btn_frame, text="إضافة/تحديث القاعدة", command=self.add_or_update_sigma_rule, style='Normal.TButton')
        add_rule_btn.pack(side=tk.LEFT, padx=5)
        remove_rule_btn = ttk.Button(btn_frame, text="حذف القاعدة المحددة", command=self.remove_sigma_rule, style='Alert.TButton')
        remove_rule_btn.pack(side=tk.LEFT, padx=5)
        self.refresh_sigma_tree()

    def refresh_sigma_tree(self):
        """تحديث قائمة قواعد SIGMA في واجهة المستخدم"""

        self.load_sigma_rules()
        for item in self.sigma_tree.get_children():
            self.sigma_tree.delete(item)
        for key, data in self.sigma_rules.items():
            self.sigma_tree.insert('', 'end', iid=key, values=(key, data['description'], data['mitre']))

    def add_or_update_sigma_rule(self):
        key = self.rule_key_entry.get().strip()
        desc = self.rule_desc_entry.get().strip()
        mitre = self.rule_mitre_entry.get().strip()
        if not key or not desc or not mitre:
            messagebox.showwarning("تنبيه", "يجب ملء جميع الحقول.")
            return
        try:
            self.c.execute("INSERT OR REPLACE INTO sigma_rules (rule_key, description, mitre) VALUES (?, ?, ?)",
                           (key, desc, mitre))
            self.conn.commit()
            self.sigma_rules[key] = {"description": desc, "mitre": mitre}
            self.refresh_sigma_tree()
            messagebox.showinfo("نجاح", "تمت إضافة/تحديث القاعدة بنجاح.")
            self.rule_key_entry.delete(0, tk.END)
            self.rule_desc_entry.delete(0, tk.END)
            self.rule_mitre_entry.delete(0, tk.END)
        except Exception as e:
            self.show_error("Sigma Update Error", e)

    def remove_sigma_rule(self):
        selected = self.sigma_tree.selection()
        if not selected:
            messagebox.showwarning("تنبيه", "لم يتم تحديد قاعدة للحذف.")
            return
        for key in selected:
            try:
                self.c.execute("DELETE FROM sigma_rules WHERE rule_key = ?", (key,))
                self.conn.commit()
                if key in self.sigma_rules:
                    del self.sigma_rules[key]
            except Exception as e:
                logging.error(f"Error removing sigma rule {key}: {str(e)}")
        self.refresh_sigma_tree()
        messagebox.showinfo("نجاح", "تم حذف القاعدة المحددة.")

    def import_sigma_rules(self):
        """استيراد قواعد SIGMA من ملف YAML أو عدة ملفات"""
        files = filedialog.askopenfilenames(title="اختر ملفات YAML", filetypes=[("YAML Files", "*.yml *.yaml")])
        if not files:
            return
        imported_count = 0
        for file in files:
            try:

                with open(file, 'r', encoding='utf-8') as f:
                    data = yaml.safe_load(f)
                    # نفترض أن بيانات YAML على شكل dict من القواعد أو قائمة من القواعد
                    if isinstance(data, dict):
                        for key, rule in data.items():
                            if isinstance(rule, dict) and "description" in rule and "mitre" in rule:
                                self.c.execute(
                                    "INSERT OR REPLACE INTO sigma_rules (rule_key, description, mitre) VALUES (?, ?, ?)",
                                    (key, rule["description"], rule["mitre"]))
                                imported_count += 1
                    elif isinstance(data, list):
                        for rule in data:
                            if isinstance(rule, dict) and "rule_key" in rule and "description" in rule and "mitre" in rule:
                                key = rule["rule_key"]
                                self.c.execute(
                                    "INSERT OR REPLACE INTO sigma_rules (rule_key, description, mitre) VALUES (?, ?, ?)",
                                    (key, rule["description"], rule["mitre"]))
                                imported_count += 1
                self.conn.commit()
            except Exception as e:
                logging.error(f"Error importing sigma rules from {file}: {str(e)}")
        self.refresh_sigma_tree()
        messagebox.showinfo("استيراد", f"تم استيراد {imported_count} قاعدة بنجاح.")

    def safe_update(self):
        try:
            self.update_all()
            self.update_dynamic_threat_intel()  # تحديث Threat Intelligence الديناميكي
        except Exception as e:
            logging.error(f"Update Error: {str(e)}")
        finally:
            self.root.after(10000, self.safe_update)

    def update_all(self):
        self.update_live_stats()
        self.populate_process_list()
        self.populate_comprehensive_list()
        self.update_pie_chart()
        self.update_comprehensive_chart()
        self.populate_history()
        self.populate_anomaly_list()

    # --- Comprehensive Scan Tab ---
    def create_comprehensive_tab(self, parent):
        toolbar = ttk.Frame(parent)
        toolbar.pack(fill=tk.X, padx=10, pady=5)
        refresh_btn = ttk.Button(toolbar, text="Refresh", command=self.update_all, style='Normal.TButton')
        refresh_btn.pack(side=tk.LEFT, padx=5)
        list_frame = ttk.Frame(parent)

        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        vsb = ttk.Scrollbar(list_frame, orient="vertical")
        hsb = ttk.Scrollbar(list_frame, orient="horizontal")
        columns = ('pid', 'name', 'path', 'cpu', 'memory', 'status', 'score')
        self.comprehensive_tree = ttk.Treeview(list_frame, columns=columns, show='headings',
                                               yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self.comprehensive_tree.yview)
        hsb.config(command=self.comprehensive_tree.xview)
        headers = [
            ('pid', 'PID', 80),
            ('name', 'Process Name', 200),
            ('path', 'Path', 350),
            ('cpu', 'CPU (%)', 100),
            ('memory', 'Memory (%)', 100),
            ('status', 'Status', 100),
            ('score', 'Suspicion Score', 120)
        ]
        for col, text, width in headers:
            self.comprehensive_tree.heading(col, text=text)
            self.comprehensive_tree.column(col, width=width, anchor=tk.CENTER)
        self.comprehensive_tree.pack(fill=tk.BOTH, expand=True)
        chart_frame = ttk.Frame(parent)
        chart_frame.pack(fill=tk.BOTH, expand=False, padx=10, pady=5)
        self.comp_figure, self.comp_ax = plt.subplots(figsize=(4, 4))
        self.comp_chart_canvas = FigureCanvasTkAgg(self.comp_figure, master=chart_frame)
        self.comp_chart_canvas.get_tk_widget().pack()
        self.comp_danger_label = ttk.Label(chart_frame, text="Danger Level: 0%", font=("Segoe UI", 12, "bold"))
        self.comp_danger_label.pack(pady=5)

    # --- Process Monitor Tab ---
    def create_monitor_tab(self, parent):
        toolbar = ttk.Frame(parent)
        toolbar.pack(fill=tk.X, padx=10, pady=5)
        refresh_btn = ttk.Button(toolbar, text="Refresh", command=self.update_all, style='Normal.TButton')
        refresh_btn.pack(side=tk.LEFT, padx=5)
        exit_btn = ttk.Button(toolbar, text="Exit", command=self.root.quit, style='Alert.TButton')
        exit_btn.pack(side=tk.RIGHT, padx=5)
        stats_frame = ttk.LabelFrame(parent, text="Live Statistics", style='Stats.TLabelframe')
        stats_frame.pack(fill=tk.X, padx=10, pady=5)
        stats_subframe = ttk.Frame(stats_frame)
        stats_subframe.pack(pady=3)
        self.cpu_label = ttk.Label(stats_subframe, text="CPU Usage: 0%", style='StatsLabel.TLabel')
        self.cpu_label.pack(side=tk.LEFT, padx=15)
        self.mem_label = ttk.Label(stats_subframe, text="Memory Usage: 0%", style='StatsLabel.TLabel')
        self.mem_label.pack(side=tk.LEFT, padx=15)
        self.process_count = ttk.Label(stats_subframe, text="Active Processes: 0", style='StatsLabel.TLabel')
        self.process_count.pack(side=tk.LEFT, padx=15)
        list_frame = ttk.Frame(parent)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        vsb = ttk.Scrollbar(list_frame, orient="vertical")

        hsb = ttk.Scrollbar(list_frame, orient="horizontal")
        columns = ('pid', 'name', 'path', 'cpu', 'memory', 'status')
        self.process_tree = ttk.Treeview(list_frame, columns=columns, show='headings',
                                         yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self.process_tree.yview)
        hsb.config(command=self.process_tree.xview)
        headers = [
            ('pid', 'PID', 80),
            ('name', 'Process Name', 200),
            ('path', 'Path', 350),
            ('cpu', 'CPU (%)', 100),
            ('memory', 'Memory (%)', 100),
            ('status', 'Status', 100)
        ]
        for col, text, width in headers:
            self.process_tree.heading(col, text=text)
            self.process_tree.column(col, width=width, anchor=tk.CENTER)
        self.process_tree.pack(fill=tk.BOTH, expand=True)
        self.process_tree.bind("<Button-3>", self.show_process_menu)
        chart_frame = ttk.Frame(parent)
        chart_frame.pack(fill=tk.BOTH, expand=False, padx=10, pady=5)
        self.figure, self.ax = plt.subplots(figsize=(4, 4))
        self.chart_canvas = FigureCanvasTkAgg(self.figure, master=chart_frame)
        self.chart_canvas.get_tk_widget().pack()
        self.danger_label = ttk.Label(chart_frame, text="Danger Level: 0%", font=("Segoe UI", 12, "bold"))
        self.danger_label.pack(pady=5)

    # --- Keyword Scan Tab ---
    def create_keyword_tab(self, parent):
        top_frame = ttk.Frame(parent)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(top_frame, text="Suspicious Keywords:").pack(side=tk.LEFT, padx=5)
        self.keyword_entry = ttk.Entry(top_frame)
        self.keyword_entry.pack(side=tk.LEFT, padx=5)
        add_btn = ttk.Button(top_frame, text="Add Keyword", command=self.add_keyword, style='Normal.TButton')
        add_btn.pack(side=tk.LEFT, padx=5)
        remove_btn = ttk.Button(top_frame, text="Remove Selected", command=self.remove_keyword, style='Alert.TButton')
        remove_btn.pack(side=tk.LEFT, padx=5)
        export_kw_btn = ttk.Button(top_frame, text="Export Keywords", command=self.export_keywords, style='Normal.TButton')
        export_kw_btn.pack(side=tk.LEFT, padx=5)
        listbox_frame = ttk.Frame(parent)
        listbox_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        self.keyword_listbox = tk.Listbox(listbox_frame, height=6)
        self.keyword_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        for kw in self.suspicious_keywords:
            self.keyword_listbox.insert(tk.END, kw)
        scrollbar = ttk.Scrollbar(listbox_frame, orient="vertical", command=self.keyword_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.keyword_listbox.config(yscrollcommand=scrollbar.set)
        scan_btn = ttk.Button(parent, text="Scan by Keywords", command=self.scan_by_keywords, style='Normal.TButton')
        scan_btn.pack(pady=5)
        result_frame = ttk.Frame(parent)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        columns = ('pid', 'name', 'path', 'match')
        self.keyword_tree = ttk.Treeview(result_frame, columns=columns, show='headings')
        for col, text, width in [('pid', 'PID', 80), ('name', 'Process Name', 200), ('path', 'Path', 350),
                                 ('match', 'Keyword Match', 150)]:
            self.keyword_tree.heading(col, text=text)
            self.keyword_tree.column(col, width=width, anchor=tk.CENTER)
        self.keyword_tree.pack(fill=tk.BOTH, expand=True)

    # --- History Tab ---
    def create_history_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        columns = ('name', 'timestamp', 'pid', 'memory', 'path', 'age', 'user', 'cpu', 'status', 'score', 'signature', 'reason', 'repetition')
        self.history_tree = ttk.Treeview(frame, columns=columns, show='tree headings')
        self.history_tree.heading("#0", text="Check")
        self.history_tree.column("#0", width=80, anchor=tk.CENTER)
        headers = [
            ('name', 'Process Name'),
            ('timestamp', 'Timestamp'),
            ('pid', 'PID'),
            ('memory', 'Memory (%)'),
            ('path', 'Path'),
            ('age', 'Age'),
            ('user', 'User'),
            ('cpu', 'CPU (%)'),
            ('status', 'Status'),
            ('score', 'Suspicion Score'),
            ('signature', 'Digital Signature'),
            ('reason', 'Reason'),
            ('repetition', 'Repetition')
        ]
        for col, text in headers:
            self.history_tree.heading(col, text=text)
            self.history_tree.column(col, width=120, anchor=tk.CENTER)
        self.history_tree.pack(fill=tk.BOTH, expand=True)
        self.history_checks = {}
        self.history_tree.bind("<Button-1>", self.on_history_click)
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        export_csv_btn = ttk.Button(btn_frame, text="Export CSV", command=self.export_history_csv, style='Normal.TButton')
        export_csv_btn.pack(side=tk.LEFT, padx=5)
        export_pdf_btn = ttk.Button(btn_frame, text="Export PDF", command=self.export_history_pdf, style='Normal.TButton')

        export_pdf_btn.pack(side=tk.LEFT, padx=5)
        export_word_btn = ttk.Button(btn_frame, text="Export Word", command=self.export_history_word, style='Normal.TButton')
        export_word_btn.pack(side=tk.LEFT, padx=5)
        export_html_btn = ttk.Button(btn_frame, text="Export HTML", command=self.export_history_html, style='Normal.TButton')
        export_html_btn.pack(side=tk.LEFT, padx=5)
        isolate_btn = ttk.Button(btn_frame, text="Isolate", command=self.isolate_suspicious, style='Alert.TButton')
        isolate_btn.pack(side=tk.LEFT, padx=5)
        block_btn = ttk.Button(btn_frame, text="Block", command=self.block_suspicious, style='Alert.TButton')
        block_btn.pack(side=tk.LEFT, padx=5)
        stop_btn = ttk.Button(btn_frame, text="Stop", command=self.stop_suspicious, style='Alert.TButton')
        stop_btn.pack(side=tk.LEFT, padx=5)
        select_all_btn = ttk.Button(btn_frame, text="Select All", command=self.select_all_history, style='Normal.TButton')
        select_all_btn.pack(side=tk.LEFT, padx=5)
        delete_selected_btn = ttk.Button(btn_frame, text="Delete Selected", command=self.delete_selected_history, style='Alert.TButton')
        delete_selected_btn.pack(side=tk.LEFT, padx=5)

    # --- Anomaly Detection Tab ---
    def create_anomaly_tab(self, parent):
        toolbar = ttk.Frame(parent)
        toolbar.pack(fill=tk.X, padx=10, pady=5)
        refresh_btn = ttk.Button(toolbar, text="Refresh", command=self.update_all, style='Normal.TButton')
        refresh_btn.pack(side=tk.LEFT, padx=5)
        list_frame = ttk.Frame(parent)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        vsb = ttk.Scrollbar(list_frame, orient="vertical")
        hsb = ttk.Scrollbar(list_frame, orient="horizontal")
        columns = ('pid', 'name', 'cpu', 'memory', 'score')
        self.anomaly_tree = ttk.Treeview(list_frame, columns=columns, show='headings',
                                         yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self.anomaly_tree.yview)
        hsb.config(command=self.anomaly_tree.xview)
        headers = [
            ('pid', 'PID', 80),
            ('name', 'Process Name', 200),
            ('cpu', 'CPU (%)', 100),
            ('memory', 'Memory (%)', 100),
            ('score', 'Anomaly Score', 120)
        ]
        for col, text, width in headers:
            self.anomaly_tree.heading(col, text=text)
            self.anomaly_tree.column(col, width=width, anchor=tk.CENTER)
        self.anomaly_tree.tag_configure("high_anomaly", background="#ffcccc")
        self.anomaly_tree.pack(fill=tk.BOTH, expand=True)

    def populate_anomaly_list(self):

        procs = []
        for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent']):
            try:
                procs.append(proc.info)
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        if not procs:
            return

        cpu_values = [p.get('cpu_percent', 0) for p in procs]
        mem_values = [p.get('memory_percent', 0) for p in procs]
        mean_cpu = sum(cpu_values) / len(cpu_values)
        mean_mem = sum(mem_values) / len(mem_values)

        def std_dev(values, mean):
            return math.sqrt(sum((x - mean)**2 for x in values) / len(values)) if len(values) > 1 else 0.01

        std_cpu = std_dev(cpu_values, mean_cpu)
        std_mem = std_dev(mem_values, mean_mem)

        anomaly_data = []
        for p in procs:
            cpu = p.get('cpu_percent', 0)
            mem = p.get('memory_percent', 0)
            z_cpu = (cpu - mean_cpu) / std_cpu if std_cpu else 0
            z_mem = (mem - mean_mem) / std_mem if std_mem else 0
            score = math.sqrt(z_cpu**2 + z_mem**2)
            anomaly_data.append((p.get('pid'), p.get('name'), cpu, mem, score))
        anomaly_data.sort(key=lambda x: x[4], reverse=True)
        self.anomaly_tree.delete(*self.anomaly_tree.get_children())
        for row in anomaly_data:
            pid, name, cpu, mem, score = row
            tags = ()
            if score >= ANOMALY_THRESHOLD:
                tags = ("high_anomaly",)
            self.anomaly_tree.insert('', 'end', values=(pid, name, f"{cpu:.1f}", f"{mem:.1f}", f"{score:.2f}"), tags=tags)

    def scan_by_keywords(self):
        try:
            self.keyword_tree.delete(*self.keyword_tree.get_children())
            for proc in psutil.process_iter(['pid', 'name', 'exe', 'cmdline']):
                try:
                    name = proc.info['name']
                    cmdline = " ".join(proc.info.get('cmdline', []))
                    match_keyword = ""
                    for keyword in self.suspicious_keywords:
                        if keyword.lower() in (name.lower() + " " + cmdline.lower()):
                            match_keyword = keyword
                            break
                    if match_keyword:

                        self.keyword_tree.insert('', 'end', values=(
                            proc.info['pid'], name, proc.info.get('exe', "N/A"), match_keyword))
                except:
                    continue
        except Exception as e:
            self.show_error("Keyword Scan Error", e)

    def add_keyword(self):
        keyword = self.keyword_entry.get().strip()
        if keyword and keyword not in self.suspicious_keywords:
            self.suspicious_keywords.append(keyword)
            self.keyword_listbox.insert(tk.END, keyword)
            self.keyword_entry.delete(0, tk.END)
            try:
                self.c.execute("INSERT OR IGNORE INTO keywords (keyword) VALUES (?)", (keyword,))
                self.conn.commit()
            except Exception as e:
                logging.error(f"Error saving keyword: {str(e)}")

    def remove_keyword(self):
        selection = self.keyword_listbox.curselection()
        if selection:
            index = selection[0]
            keyword = self.keyword_listbox.get(index)
            self.suspicious_keywords.remove(keyword)
            self.keyword_listbox.delete(index)
            try:
                self.c.execute("DELETE FROM keywords WHERE keyword=?", (keyword,))
                self.conn.commit()
            except Exception as e:
                logging.error(f"Error removing keyword: {str(e)}")

    def export_keywords(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".csv",
                                                     filetypes=[("CSV files", "*.csv")],
                                                     title="Save Keywords CSV")
            if not file_path:
                return
            self.c.execute("SELECT keyword FROM keywords")
            rows = self.c.fetchall()
            with open(file_path, mode='w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerow(["Keyword"])
                for row in rows:
                    writer.writerow(row)
            messagebox.showinfo("Success", "Keywords exported successfully.")
        except Exception as e:
            self.show_error("Export Keywords Error", e)

    def populate_history(self):

        try:
            self.history_tree.delete(*self.history_tree.get_children())
            self.c.execute("SELECT id, timestamp, process_name, pid, cpu_percent, memory_percent, path, command_line, status, suspicion_score, reason, user_info FROM process_history ORDER BY id DESC")
            rows = self.c.fetchall()
            today = datetime.now().date()
            history_dict = {}
            for row in rows:
                record_id, timestamp_str, name, pid, cpu_usage, mem_usage, path, cmdline, status, score, reason, user_info = row
                timestamp_dt = datetime.strptime(timestamp_str, "%Y-%m-%d %H:%M:%S")
                age = str(datetime.now() - timestamp_dt).split('.')[0]
                signature = hashlib.md5(cmdline.encode()).hexdigest() if cmdline != "N/A" else "N/A"
                key = (name, today)
                history_dict[key] = history_dict.get(key, 0) + 1
                repetition = history_dict[key]
                self.history_tree.insert('', 'end', iid=str(record_id), text='',
                    values=(name, timestamp_str, pid, f"{mem_usage:.1f}", path, age, user_info, f"{cpu_usage:.1f}", status, f"{score:.2f}", signature, reason, repetition)
                )
                self.history_checks[str(record_id)] = True
                self.history_tree.item(str(record_id), image=self.get_checkbox_image(True))
        except Exception as e:
            self.show_error("Error populating history", e)

    def get_checkbox_image(self, checked):
        return self.checkbox_checked if checked else self.checkbox_unchecked

    def on_history_click(self, event):
        region = self.history_tree.identify("region", event.x, event.y)
        if region == "tree":
            item = self.history_tree.identify_row(event.y)
            if item:
                current = self.history_checks.get(item, True)
                new_state = not current
                self.history_checks[item] = new_state
                self.history_tree.item(item, image=self.get_checkbox_image(new_state))

    def export_history_csv(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".csv",
                                                     filetypes=[("CSV files", "*.csv")],
                                                     title="Save History CSV")
            if not file_path:
                return
            headers = ["Process Name", "Timestamp", "PID", "Memory (%)", "Path", "Age", "User", "CPU (%)", "Status",
                       "Suspicion Score", "Digital Signature", "Reason", "Repetition"]
            with open(file_path, mode='w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerow(headers)

                for item in self.history_tree.get_children():
                    values = self.history_tree.item(item, 'values')
                    writer.writerow(values)
            messagebox.showinfo("Success", "History exported successfully.")
        except Exception as e:
            self.show_error("Export CSV Error", e)

    def export_history_pdf(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".pdf",
                                                     filetypes=[("PDF files", "*.pdf")],
                                                     title="Save History PDF")
            if not file_path:
                return
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter
            y = height - 50
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, "History Report")
            y -= 30
            c.setFont("Helvetica", 10)
            headers = ["Process Name", "Timestamp", "PID", "Memory (%)", "Path", "Age", "User", "CPU (%)", "Status",
                       "Suspicion Score", "Digital Signature", "Reason", "Repetition"]
            c.drawString(50, y, " | ".join(headers))
            y -= 20
            for item in self.history_tree.get_children():
                values = self.history_tree.item(item, 'values')
                c.drawString(50, y, " | ".join(str(v) for v in values))
                y -= 15
                if y < 50:
                    c.showPage()
                    y = height - 50
            c.save()
            messagebox.showinfo("Success", "History exported successfully.")
        except Exception as e:
            self.show_error("Export PDF Error", e)

    def export_history_word(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")],
                                                     title="Save History Word")
            if not file_path:
                return
            from docx import Document
            document = Document()
            document.add_heading("History Report", level=1)
            headers = ["Process Name", "Timestamp", "PID", "Memory (%)", "Path", "Age", "User", "CPU (%)",

"Status",
                       "Suspicion Score", "Digital Signature", "Reason", "Repetition"]
            table = document.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            for item in self.history_tree.get_children():
                values = self.history_tree.item(item, 'values')
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
            document.save(file_path)
            messagebox.showinfo("Success", "History exported successfully.")
        except Exception as e:
            self.show_error("Export Word Error", e)

    def export_history_html(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".html",
                                                     filetypes=[("HTML files", "*.html")],
                                                     title="Save History HTML")
            if not file_path:
                return
            headers = ["Process Name", "Timestamp", "PID", "Memory (%)", "Path", "Age", "User", "CPU (%)", "Status",
                       "Suspicion Score", "Digital Signature", "Reason", "Repetition"]
            html = "<html><head><meta charset='utf-8'><title>History Report</title>"
            html += "<style>table { border-collapse: collapse; width: 100%; } th, td { border: 1px solid #ddd; padding: 8px; } th { background-color: #2c3e50; color: white; }</style>"
            html += "</head><body>"
            html += "<h1>History Report</h1><table><tr>"
            for header in headers:
                html += f"<th>{header}</th>"
            html += "</tr>"
            for item in self.history_tree.get_children():
                values = self.history_tree.item(item, 'values')
                html += "<tr>" + "".join(f"<td>{v}</td>" for v in values) + "</tr>"
            html += "</table></body></html>"
            with open(file_path, "w", encoding='utf-8') as f:
                f.write(html)
            messagebox.showinfo("Success", "History exported as HTML successfully.")
        except Exception as e:
            self.show_error("Export HTML Error", e)

    def select_all_history(self):
        for item in self.history_tree.get_children():
            self.history_checks[item] = True
            self.history_tree.item(item, image=self.get_checkbox_image(True))

    def delete_selected_history(self):
        to_delete = []

        for item, checked in self.history_checks.items():
            if checked:
                to_delete.append(item)
        if not to_delete:
            messagebox.showinfo("Delete", "لم يتم تحديد أي بيانات للحذف.")
            return
        if not messagebox.askyesno("Delete", "هل أنت متأكد من حذف البيانات المحددة؟"):
            return
        for item in to_delete:
            try:
                record_id = int(item)
                self.c.execute("DELETE FROM process_history WHERE id = ?", (record_id,))
                self.conn.commit()
                self.history_tree.delete(item)
                del self.history_checks[item]
            except Exception as e:
                logging.error(f"Error deleting history item {item}: {str(e)}")
        messagebox.showinfo("Delete", "تم حذف البيانات المحددة.")

    def isolate_suspicious(self):
        messagebox.showinfo("Action", "تم تنفيذ عملية العزل (محاكاة).")

    def block_suspicious(self):
        messagebox.showinfo("Action", "تم تنفيذ عملية الحظر (محاكاة).")

    def stop_suspicious(self):
        messagebox.showinfo("Action", "تم تنفيذ عملية الإيقاف (محاكاة).")

    def show_process_menu(self, event):
        item = self.process_tree.identify_row(event.y)
        if item:
            self.process_tree.selection_set(item)
            menu = tk.Menu(self.root, tearoff=0, bg="#c2e9fb", activebackground="#a1c4fd", font=("Segoe UI", 10))
            menu.add_command(label="إيقاف مؤقت", command=self.suspend_selected_process)
            menu.add_command(label="إنهاء", command=self.terminate_selected_process)
            menu.tk_popup(event.x_root, event.y_root)

    def suspend_selected_process(self):
        selected = self.process_tree.selection()
        if selected:
            pid = int(self.process_tree.item(selected[0], 'values')[0])
            try:
                proc = psutil.Process(pid)
                proc.suspend()
                messagebox.showinfo("Process Suspended", f"تم إيقاف العملية {pid} مؤقتًا.")
            except Exception as e:
                self.show_error("Suspend Error", e)

    def terminate_selected_process(self):
        selected = self.process_tree.selection()

        if selected:
            pid = int(self.process_tree.item(selected[0], 'values')[0])
            try:
                proc = psutil.Process(pid)
                proc.terminate()
                messagebox.showinfo("Process Terminated", f"تم إنهاء العملية {pid}.")
            except Exception as e:
                self.show_error("Terminate Error", e)

    def show_error(self, title, error):
        logging.error(f"{title}: {str(error)}")
        messagebox.showerror(title, f"{str(error)}")

    # --- تعديل دالة حساب درجة الاشتباه لدمج قواعد SIGMA وMitre ATT&CK ---
    def calculate_suspicion_score(self, proc_info):
        score = 0.0
        reasons = []
        name = proc_info.get('name', '').lower()
        cmdline = " ".join(proc_info.get('cmdline', [])).lower()
        cpu_usage = proc_info.get('cpu_percent', 0.0)
        mem_usage = proc_info.get('memory_percent', 0.0)

        # التحقق من أسماء العمليات المشبوهة
        if name in [n.lower() for n in self.threat_intel.get("suspicious_names", [])]:
            score += 0.3
            reasons.append("Matched static threat intel (suspicious name)")
        # التحقق من الكلمات المفتاحية الثابتة
        for keyword in self.suspicious_keywords:
            if keyword.lower() in (name + " " + cmdline):
                score += 0.3
                reasons.append(f"Matched keyword: {keyword}")
                break
        # التحقق من استخدام الموارد المرتفعة
        if cpu_usage >= HIGH_RESOURCE_THRESHOLD['cpu']:
            score += 0.2
            reasons.append("High CPU usage")
        if mem_usage >= HIGH_RESOURCE_THRESHOLD['memory']:
            score += 0.2
            reasons.append("High Memory usage")
        # دمج قواعد SIGMA وإضافة معرّفات MITRE من القواعد
        sigma_bonus = 0.0
        sigma_matches = []
        for rule_key, rule_data in self.sigma_rules.items():
            if rule_key.lower() in (name + " " + cmdline):
                sigma_bonus += 0.1
                sigma_matches.append(f"{rule_key} (MITRE {rule_data['mitre']})")
        if sigma_matches:
            reasons.append("Sigma/MITRE rules matched: " + ", ".join(sigma_matches))
        total_score = min(score + sigma_bonus, 1.0)
        proc_info["reason_detail"] = "; ".join(reasons) if reasons else ""
        return total_score


    def evaluate_process_ml(self, proc_info):
        """استخراج الميزات وتمريرها للنموذج لتقييم العملية (0: طبيعي، 1: مشبوهة)"""
        if self.ml_model is None:
            return 0
        features = self.extract_features(proc_info)
        prediction = self.ml_model.predict(features)
        return int(prediction[0])

    def populate_comprehensive_list(self):
        try:
            self.comprehensive_tree.delete(*self.comprehensive_tree.get_children())
            for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'status', 'exe', 'cmdline']):
                try:
                    path = proc.info.get('exe', "N/A")
                    pid = proc.info['pid']
                    name = proc.info['name']
                    cpu_usage = proc.info.get('cpu_percent', 0.0)
                    mem_usage = proc.info.get('memory_percent', 0.0)
                    status = proc.info.get('status', "Unknown")
                    cmdline = " ".join(proc.info.get('cmdline', []))
                    static_score = self.calculate_suspicion_score(proc.info)
                    ml_label = self.evaluate_process_ml(proc.info)
                    total_score = static_score + (0.5 if ml_label == 1 else 0)
                    total_score = min(total_score, 1.0)
                    tag = "suspicious" if total_score > 0.5 else ""
                    self.comprehensive_tree.insert('', 'end',
                        values=(pid, name, path, f"{cpu_usage:.1f}", f"{mem_usage:.1f}", status, f"{total_score:.2f}"),
                        tags=(tag,))
                    if tag:
                        self.comprehensive_tree.tag_configure("suspicious", background=COLORS['danger'])
                    if total_score > 0.5:
                        reason = proc.info.get("reason_detail", "Suspicious Process")
                        self.save_history(name, pid, cpu_usage, mem_usage, path, cmdline, status, reason)
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
                except Exception as e:
                    logging.error(f"Error in comprehensive list: {str(e)}")
        except Exception as e:
            self.show_error("Error populating comprehensive list", e)

    def populate_process_list(self):
        try:
            self.process_tree.delete(*self.process_tree.get_children())
            for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'status', 'exe', 'cmdline']):
                try:
                    cpu_usage = proc.info.get('cpu_percent', 0.0)
                    mem_usage = proc.info.get('memory_percent', 0.0)
                    path = proc.info.get('exe', "N/A")

                    pid = proc.info['pid']
                    name = proc.info['name']
                    status = proc.info.get('status', "Unknown")
                    tag = ""
                    if cpu_usage >= HIGH_RESOURCE_THRESHOLD['cpu'] or mem_usage >= HIGH_RESOURCE_THRESHOLD['memory']:
                        tag = "high_usage"
                    self.process_tree.insert('', 'end',
                        values=(pid, name, path, f"{cpu_usage:.1f}", f"{mem_usage:.1f}", status),
                        tags=(tag,))
                    if tag:
                        self.process_tree.tag_configure("high_usage", background=COLORS['high_usage'])
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
                except Exception as e:
                    logging.error(f"Error in process list: {str(e)}")
        except Exception as e:
            self.show_error("Error populating process list", e)

    def update_pie_chart(self):
        try:
            total = len(psutil.pids())
            suspicious_count = 0
            for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'exe', 'cmdline']):
                try:
                    name = proc.info['name']
                    cmdline = " ".join(proc.info.get('cmdline', []))
                    is_suspicious = False
                    if name.lower() in [n.lower() for n in self.threat_intel.get("suspicious_names", [])]:
                        is_suspicious = True
                    for keyword in self.suspicious_keywords:
                        if keyword.lower() in (name.lower() + " " + cmdline.lower()):
                            is_suspicious = True
                            break
                    if is_suspicious:
                        suspicious_count += 1
                except:
                    continue
            danger_level = (suspicious_count / total * 100) if total > 0 else 0
            self.danger_label.config(text=f"Danger Level: {danger_level:.1f}%")
            self.ax.clear()
            self.ax.pie([suspicious_count, total - suspicious_count],
                        labels=["Suspicious", "Normal"],
                        colors=[COLORS['danger'], COLORS['safe']],
                        autopct='%1.1f%%')
            self.figure.tight_layout()
            self.chart_canvas.draw()
        except Exception as e:
            logging.error(f"Error updating pie chart: {str(e)}")

    def update_comprehensive_chart(self):

        try:
            total = len(psutil.pids())
            suspicious_count = 0
            for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'exe', 'cmdline']):
                try:
                    name = proc.info['name']
                    cmdline = " ".join(proc.info.get('cmdline', []))
                    is_suspicious = False
                    if name.lower() in [n.lower() for n in self.threat_intel.get("suspicious_names", [])]:
                        is_suspicious = True
                    for keyword in self.suspicious_keywords:
                        if keyword.lower() in (name.lower() + " " + cmdline.lower()):
                            is_suspicious = True
                            break
                    if is_suspicious:
                        suspicious_count += 1
                except:
                    continue
            danger_level = (suspicious_count / total * 100) if total > 0 else 0
            self.comp_danger_label.config(text=f"Danger Level: {danger_level:.1f}%")
            self.comp_ax.clear()
            self.comp_ax.pie([suspicious_count, total - suspicious_count],
                             labels=["Suspicious", "Normal"],
                             colors=[COLORS['danger'], COLORS['safe']],
                             autopct='%1.1f%%')
            self.comp_figure.tight_layout()
            self.comp_chart_canvas.draw()
        except Exception as e:
            logging.error(f"Error updating comprehensive chart: {str(e)}")

    def save_history(self, name, pid, cpu_usage, mem_usage, path, cmdline, status, reason):
        try:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            username = getpass.getuser()
            user_info = f"{username} (Admin)" if is_admin() else f"{username} (User)"
            proc_info = {
                "name": name,
                "cmdline": cmdline.split(),
                "cpu_percent": cpu_usage,
                "memory_percent": mem_usage
            }
            score = self.calculate_suspicion_score(proc_info)
            detailed_reason = proc_info.get("reason_detail", reason)
            self.c.execute(
                "INSERT INTO process_history (timestamp, process_name, pid, cpu_percent, memory_percent, path, command_line, status, suspicion_score, reason, user_info) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (timestamp, name, pid, cpu_usage, mem_usage, path, cmdline, status, score, detailed_reason, user_info)
            )
            self.conn.commit()
        except Exception as e:

            logging.error(f"Error saving history: {str(e)}")

    def update_live_stats(self):
        try:
            cpu = psutil.cpu_percent()
            mem = psutil.virtual_memory().percent
            count = len(psutil.pids())
            self.cpu_label.config(text=f"CPU Usage: {cpu}%")
            self.mem_label.config(text=f"Memory Usage: {mem}%")
            self.process_count.config(text=f"Active Processes: {count}")
        except Exception as e:
            logging.error(f"Error updating stats: {str(e)}")

if __name__ == "__main__":
    if not is_admin():
        tk.messagebox.showerror("Permission Error", "يجب تشغيل هذا البرنامج بصلاحيات المدير!")
        sys.exit(1)
    try:
        initialize_database()
        root = tk.Tk()
        app = IntegratedTaskManager(root)
        root.mainloop()
    except Exception as e:
        logging.critical(f"Critical error: {str(e)}", exc_info=True)
        tk.messagebox.showerror("Critical Error", f"حدث خطأ غير متوقع: {str(e)}\nيرجى مراجعة سجل الأخطاء.")
        sys.exit(1)















































