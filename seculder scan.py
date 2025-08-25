#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
NTST.py -- Scheduled Tasks Scanner (نسخة محسنة بعد معالجة المشكلات والتحسينات المطلوبة)
التعديلات:
1- استبدال حجم العرض الافتراضي بوضع ملء الشاشة وعدم استخدام الحجم الثابت.
2- تصحيح زر التصدير بحيث لا ينتج عنه خطأ عند التصدير وتعديل لونه ليكون مثل باقي أزرار التطبيق ويمتزج مع الثيمات.
3- تعديل تخطيط العرض في قسم المدخلات والإحصاءات لجعلهما متوازيين، حيث يتم تكبير قسم المدخلات وتصغير قسم الإحصاءات ليصبحا متساويين بشكل جمالي.
4- إزالة زر "استعادة خيارات التوقيع" نظراً لتعذر إظهاره بوظيفة مفيدة.
5- تفعيل مؤشر "توقيع تنفيذي غير صالح/مفقود" في الرسم البياني للإحصاءات.
6- تعديل نص مؤشر مطابقة الأمر: استبدال النص "مطابقة كلمة (وسائط/أمر)" بالنص "العوامل (Arguments)" في قسم الإحصاءات.
7- استبدال النص في قائمة النتائج في عمود "سبب الاشتباه" عند مطابقة الشرط؛ إذ إذا كان السبب يبدأ بـ "العوامل (Arguments)" يتم تغييره إلى "مطابقة للعوامل/ كلمات مفتاحية" مع بقاء رسالة السبب.
8- ضمان تطبيق هذه التعديلات مع دعم اللغتين العربية والإنجليزية.
9- إزالة التصدير كملف PDF.
10- تحسين عرض الواجهة بشكل عام.
11- تحسين الخط في زر التصدير لجعله أكثر وضوحاً.
12- إضافة خيار "عرض التفاصيل" عند النقر بزر الماوس الأيمن على سجل في الجدول.
"""

import sys, os, re, json, base64, traceback, subprocess, shlex, pythoncom
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Dict, Any, Optional

# اعتمادات PyQt5
try:
    from PyQt5.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QFrame,
        QPushButton, QLabel, QLineEdit, QListWidget, QListWidgetItem, QInputDialog,
        QComboBox, QSpinBox, QTableWidget, QTableWidgetItem, QHeaderView,
        QAbstractItemView, QMessageBox, QFileDialog, QStatusBar, QProgressBar, QSizePolicy, QDialog,
        QTextEdit, QMenu, QToolButton
    )
    from PyQt5.QtCore import Qt, QThread, pyqtSignal, QByteArray, QEvent, QTimer, QPoint
    from PyQt5.QtGui import QPixmap, QKeySequence
except Exception as e:
    raise ImportError("PyQt5 مطلوب: pip install PyQt5") from e

# اختيارية: رسم بياني (pyqtgraph)
try:
    import pyqtgraph as pg
except Exception:
    pg = None  # سيتم التعامل مع عدم توفره في الواجهة

# اختيارية: Excel (openpyxl)
try:
    import openpyxl
    from openpyxl.utils import get_column_letter
except Exception:
    openpyxl = None

# إضافات للتصدير بصيغ Word وCSV (تم إزالة PDF)
try:
    from docx import Document
except Exception:
    Document = None

# تم إزالة استيراد reportlab و canvas

# إضافات COM (win32com)
try:
    import win32com.client

    HAVE_COM = True
except Exception:
    HAVE_COM = False

# مسار التطبيق لحفظ/تحميل القوائم
APP_DIR = Path.home() / ".ntst_ui"
APP_DIR.mkdir(exist_ok=True)
LISTS_FILE = APP_DIR / "lists_tasks.json"

# ================= شعار Base64 =================
SAFE_LOGO_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAHgAAAB4CAYAAAA5ZDbfAAAABHNCSVQICAgIfAhkiAAAAAlwSFlz"
    "AAAXEgAAFxIBZ5/SUgAAABlLV0Q3JlYXRpb24gVGltZQAwOC8xNy8yMDI1w3t3YAAAcdhJFU"
    "eJzt2zFPG0ccB/Dn5o0Wg2bY2Ckq2QfJx5q1QXyQ1x0P8bqYkQmZlZKXqgQdJm+Z2eNw7cO3cM5N"
    "b3wN2dYt7bG2Q0p7b3C0Ck5mYV9dB8KQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH8l7m5m7l7f8w"
    "qj2o9yVx1H3y4o7X5q3wQ1V7b7mXw9Vv8C0d8p7yC2fQ2v5y6t6XvIYH2z9q3p9K5m3f7mXk2x3z"
    "mJz0n8Y0o4bY2u1u9kQmZ9zqYg0q7hLwQyFkH3Gflr8a2m4E1O2C7rrm0b/vN4h1n9FfQfV0r3Yx"
    "n7J3Yq1xkFmK8y5o6H2p7B2vYVx/0cU9k4o7Em8bq2lq0sRj1mL0k3x7o3n3PAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAPwHfwC0v0q+u0JQ8wAAAABJRU5ErkJggg=="
)

def pixmap_from_base64(b64: str) -> QPixmap:
    pm = QPixmap()
    try:
        if not b64:
            return pm
        pad = len(b64) % 4
        if pad:
            b64 += "=" * (4 - pad)
        data = base64.b64decode(b64, validate=False)
        try:
            pm.loadFromData(QByteArray(data))
        except Exception:
            pm.loadFromData(data)
        return pm
    except Exception:
        return QPixmap()

# ================= الترجمات =================
# تعديل نصوص مطابقة الكلمات والنصوص الخاصة بالنتائج
L_AR = {
    "title": "أداة فحص المهام المجدولة -- واجهة ملونة",
    "scan": "بدء الفحص",
    "stop": "إيقاف",
    "clear": "مسح",
    "export": "تصدير",
    "export_xlsx": "تصدير بصيغة XLSX",
    # "export_pdf": "تصدير بصيغة PDF", # تم إزالة هذا الخيار
    "export_word": "تصدير بصيغة Word",
    "export_txt": "تصدير بصيغة TXT",
    "export_csv": "تصدير بصيغة CSV",
    "exit": "خروج",
    "theme": "الثيم",
    "lang": "اللغة",
    "inputs": "المدخلات",
    "stats": "إحصاءات الفحص",
    "reasons_chart": "أسباب الاشتباه",
    "results": "نتائج الفحص",
    "filter_hint": "فلترة النتائج (Ctrl+F للإظهار/الإخفاء)",
    "kw_args": "كلمات مفتاحية/عوامل",
    "kw_filter": "فلتر الكلمات",
    "add": "+",
    "remove": "−",
    "schedule_type": "نوع الجدولة",
    "schedule_any": "الكل",
    "schedule_daily": "يومي",
    "schedule_weekly": "أسبوعي",
    "schedule_monthly": "شهري",
    "schedule_once": "مرة واحدة",
    "user": "المستخدم",
    "search_bar_placeholder": "اكتب للفلترة عبر أي عمود...",
    "need_openpyxl": "ثبّت openpyxl للتصدير بصيغة XLSX: pip install openpyxl",
    "need_docx": "ثبّت python-docx للتصدير بصيغة Word: pip install python-docx",
    # "need_reportlab": "ثبّت reportlab للتصدير بصيغة PDF: pip install reportlab", # تم إزالة هذا الخيار
    "progress": "جاري الفحص...",
    "done": "انتهى الفحص: {} مهمة مشبوهة من إجمالي {} مهمة.",
    "no_filters": "الرجاء تحديد معيار واحد على الأقل قبل الفحص (كلمة أو نوع جدولة).",
    "tbl_headers": [
        "اسم المهمة", "المسار", "الأمر", "الوسائط", "المؤلف/الحساب", "تاريخ الإنشاء", "نوع الجدولة", "التوقيع", "سبب الاشتباه"
    ],
    "reason_kw_arg": "العوامل (Arguments)",  # تم تعديل النص هنا
    "reason_kw_arg_result": "مطابقة للعوامل/ كلمات مفتاحية",  # النص الجديد لنتائج العمود
    "reason_sig": "توقيع تنفيذي غير صالح/مفقود",
    "state_ok": "وصول",
    "state_denied": "مرفوض",
    "details": "تفاصيل المهمة",
    "created": "تاريخ الإنشاء",
    "unknown": "غير معروف",
    "ok": "موافق",
    "show_details": "عرض التفاصيل" # إضافة نص جديد لقائمة السياق
}

L_EN = {
    "title": "Scheduled Tasks Scanner -- Geometric Colored UI",
    "scan": "Scan",
    "stop": "Stop",
    "clear": "Clear",
    "export": "Export",
    "export_xlsx": "Export as XLSX",
    # "export_pdf": "Export as PDF", # تم إزالة هذا الخيار
    "export_word": "Export as Word",
    "export_txt": "Export as TXT",
    "export_csv": "Export as CSV",
    "exit": "Exit",
    "theme": "Theme",
    "lang": "Language",
    "inputs": "Inputs",
    "stats": "Scan Stats",
    "reasons_chart": "Suspicion Reasons",
    "results": "Scan Results",
    "filter_hint": "Filter (Ctrl+F to toggle)",
    "kw_args": "Keywords/Factors",
    "kw_filter": "Keyword filter",
    "add": "+",
    "remove": "−",
    "schedule_type": "Schedule type",
    "schedule_any": "All",
    "schedule_daily": "Daily",
    "schedule_weekly": "Weekly",
    "schedule_monthly": "Monthly",
    "schedule_once": "Once",
    "user": "User",
    "search_bar_placeholder": "Type to filter any column...",
    "need_openpyxl": "Install openpyxl for XLSX export: pip install openpyxl",
    "need_docx": "Install python-docx for Word export: pip install python-docx",
    # "need_reportlab": "Install reportlab for PDF export: pip install reportlab", # تم إزالة هذا الخيار
    "progress": "Scanning...",
    "done": "Done: {} suspicious of {} tasks.",
    "no_filters": "Please enable at least one filter before scanning (keyword or schedule).",
    "tbl_headers": [
        "Task name", "Path", "Command", "Arguments", "Author/Account", "Created", "Schedule", "Signature", "Reasons"
    ],
    "reason_kw_arg": "Arguments",  # تم تعديل النص هنا
    "reason_kw_arg_result": "Matches Keywords/Arguments",  # النص الجديد لنتائج العمود
    "reason_sig": "Invalid/Missing signature",
    "state_ok": "Access",
    "state_denied": "Denied",
    "details": "Task details",
    "created": "Created",
    "unknown": "N/A",
    "ok": "OK",
    "show_details": "Show Details" # إضافة نص جديد لقائمة السياق
}

LANG = "ar"

def tr(key: str) -> str:
    return (L_AR if LANG == "ar" else L_EN).get(key, key)

# ================= مساعدات =================
def split_tokens(raw: List[str]) -> List[str]:
    out = []
    for s in raw:
        for part in str(s).split(','):
            t = part.strip()
            if t:
                out.append(t)
    return out

def exact_token_present(text: str, tokens: List[str]) -> Optional[str]:
    if not text or not tokens:
        return None
    low_text = text.lower()
    for t in tokens:
        pat = r'(?<![0-9A-Za-z_])' + re.escape(t.lower()) + r'(?![0-9A-Za-z_])'
        if re.search(pat, low_text, flags=re.IGNORECASE):
            return t
    return None

# ================= تمثيل المعايير =================
@dataclass
class Criteria:
    kw_args: List[str] = field(default_factory=list)
    sched_type: str = "all"
    user: str = ""

# ================= تواريخ/جدولة =================
def parse_iso_dt(s: str) -> Optional[datetime]:
    if not s:
        return None
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00")).replace(tzinfo=None)
    except Exception:
        for f in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%m/%d/%Y %H:%M:%S"):
            try:
                return datetime.strptime(s, f)
            except Exception:
                pass
    return None

def schedule_from_triggers(triggers: Any) -> str:
    try:
        seen = set()
        for t in triggers:
            typ = None
            if hasattr(t, "Type"):
                typ = int(t.Type)
            elif isinstance(t, dict):
                typ = t.get("Type")
            if typ in (1, "Once"):
                seen.add("Once")
            elif typ in (2, "Daily"):
                seen.add("Daily")
            elif typ in (3, "Weekly"):
                seen.add("Weekly")
            elif typ in (4, 5, "Monthly"):
                seen.add("Monthly")
        if not seen:
            return "All"
        return "All" if len(seen) > 1 else next(iter(seen))
    except Exception:
        return "All"

def created_from_registration(reginfo: Any) -> Optional[datetime]:
    try:
        if hasattr(reginfo, "Date"):
            return parse_iso_dt(reginfo.Date)
        if isinstance(reginfo, dict):
            return parse_iso_dt(reginfo.get("Date"))
    except Exception:
        pass
    return None

def author_from(reginfo: Any, principal: Any) -> str:
    a = ""
    try:
        if hasattr(reginfo, "Author"):
            a = reginfo.Author or ""
    except Exception:
        pass
    try:
        if not a and hasattr(principal, "UserId"):
            a = principal.UserId or ""
    except Exception:
        pass
    if isinstance(reginfo, dict):
        a = a or reginfo.get("Author", "")
    if isinstance(principal, dict):
        a = a or principal.get("UserId", "")
    return a or tr("unknown")

# ================= فحص توقيع الملف التنفيذي =================
def check_signature_status(exe_path: str) -> str:
    if not exe_path:
        return "N/A"
    try:
        if not os.path.isfile(exe_path):
            return "N/A"
        ps = os.path.join(os.environ.get("SystemRoot", r"C:\\Windows"),
                          r"System32\\WindowsPowerShell\\v1.0\\powershell.exe")
        if not os.path.isfile(ps):
            return "N/A"
        quoted_path = shlex.quote(exe_path)
        cmd = f'"{ps}" -NoProfile -ExecutionPolicy Bypass -Command "try {{ (Get-AuthenticodeSignature {quoted_path}).Status }} catch {{ \\"N/A\\" }}"'
        p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True, text=True, encoding="utf-8",
                           errors="ignore")
        out = (p.stdout or "").strip()
        if not out:
            return "N/A"
        out = out.splitlines()[-1].strip()
        if "Valid" in out:
            return "Valid"
        if "NotSigned" in out:
            return "Unsigned"
        if any(x in out for x in ("UnknownError", "HashMismatch", "NotTrusted", "Invalid")):
            return "Invalid"
        return out[:32]
    except Exception:
        return "N/A"

# ================= محول آمن للنص =================
def safe_str(x: Any) -> str:
    try:
        return str(x)
    except Exception:
        return repr(x)

# ================= ماسح المهام =================
class TasksScannerThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(list, int)
    error = pyqtSignal(str)

    def __init__(self, crit: Criteria):
        super().__init__()
        self.crit = crit
        self._stop = False

    def stop(self):
        self._stop = True

    def _match_args_keywords(self, command: str, args: str, kw_tokens: List[str]) -> Optional[str]:
        txt = " ".join([command or "", args or ""]).strip()
        return exact_token_present(txt, kw_tokens)

    def _sched_match(self, sched: str, want: str) -> bool:
        if not want or want.lower() == "all":
            return False
        return (sched or "").lower() == want.lower()

    def _collect_tasks_com(self) -> List[Dict[str, Any]]:
        srv = win32com.client.Dispatch("Schedule.Service")
        srv.Connect()
        out: List[Dict[str, Any]] = []
        self._collect_folder_com(srv.GetFolder("\\"), "\\", out)
        return out

    def _collect_folder_com(self, folder, path: str, out: List[Dict[str, Any]]):
        if self._stop:
            return
        try:
            tasks = folder.GetTasks(1)
            total = tasks.Count
        except Exception:
            tasks, total = [], 0
        for i in range(1, total + 1):
            if self._stop:
                return
            try:
                t = tasks.Item(i)
            except Exception:
                continue
            name = desc = cmd = args = author = schedule = ""
            created_dt = None
            try:
                name = safe_str(t.Name)
                d = t.Definition
                ri = getattr(d, "RegistrationInfo", None)
                author = author_from(ri, d.Principal)
                desc = safe_str(getattr(ri, "Description", "")) if ri else ""
                created_dt = created_from_registration(ri) or None
                schedule = schedule_from_triggers(d.Triggers)
                acts = d.Actions
                for ai in range(1, getattr(acts, "Count", 0) + 1):
                    a = acts.Item(ai)
                    typ = None
                    try:
                        typ = int(a.Type)
                    except Exception:
                        pass
                    if typ == 0 or hasattr(a, "Path"):
                        cmd = safe_str(getattr(a, "Path", "")).strip()
                        args = safe_str(getattr(a, "Arguments", "")).strip()
                        if cmd:
                            break
            except Exception:
                pass
            out.append({
                "name": name,
                "path": path.strip("\\"),
                "desc": desc,
                "command": cmd,
                "arguments": args,
                "author": author,
                "created": created_dt,
                "schedule": schedule
            })
            self.progress.emit(len(out))
        try:
            subs = folder.GetFolders(0)
            for j in range(1, subs.Count + 1):
                if self._stop:
                    return
                sub = subs.Item(j)
                sub_path = os.path.join(path, safe_str(sub.Name))
                self._collect_folder_com(sub, sub_path, out)
        except Exception:
            pass

    def _collect_tasks_schtasks(self) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        try:
            p_list = subprocess.run(
                r'schtasks /Query /FO LIST /V',
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True,
                text=True, encoding="utf-8", errors="ignore"
            )
            text_list = p_list.stdout
            task_names_raw = []
            cur_task_info = {}
            for line in text_list.splitlines():
                line = line.strip()
                if not line:
                    if cur_task_info and "TaskName" in cur_task_info:
                        task_names_raw.append(cur_task_info["TaskName"].strip())
                    cur_task_info = {}
                    continue
                if ":" in line:
                    k, v = line.split(":", 1)
                    cur_task_info[k.strip()] = v.strip()
            if cur_task_info and "TaskName" in cur_task_info:
                task_names_raw.append(cur_task_info["TaskName"].strip())

            for tn in task_names_raw:
                if self._stop:
                    return out
                if not tn.startswith("\\"):
                    tn = "\\" + tn
                if "\\" in tn[1:]:
                    path_part = os.path.dirname(tn)
                    name_part = os.path.basename(tn)
                else:
                    path_part = "\\"
                    name_part = tn.strip("\\")
                try:
                    px_xml = subprocess.run(
                        f'schtasks /Query /TN "{tn}" /XML',
                        stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True,
                        text=True, encoding="utf-8", errors="ignore"
                    )
                    xml_content = px_xml.stdout

                    def xfind(tag: str) -> str:
                        m = re.search(rf"<{tag}>(.*?)</{tag}>", xml_content, flags=re.I | re.S)
                        return (m.group(1).strip() if m else "")

                    desc = xfind("Description")
                    author = xfind("Author") or tr("unknown")
                    created_dt = parse_iso_dt(xfind("Date"))
                    command = xfind("Command")
                    arguments = xfind("Arguments")
                    trig_type = "All"
                    if re.search(r"<TimeTrigger", xml_content, re.I):
                        trig_type = "Once"
                    elif re.search(r"<DailyTrigger", xml_content, re.I):
                        trig_type = "Daily"
                    elif re.search(r"<WeeklyTrigger", xml_content, re.I):
                        trig_type = "Weekly"
                    elif re.search(r"<Monthly", xml_content, re.I):
                        trig_type = "Monthly"
                    schedule = trig_type
                except Exception:
                    desc = ""
                    author = tr("unknown")
                    created_dt = None
                    command = ""
                    arguments = ""
                    schedule = "All"
                out.append({
                    "name": name_part,
                    "path": path_part.strip("\\"),
                    "desc": desc,
                    "command": command,
                    "arguments": arguments,
                    "author": author,
                    "created": created_dt,
                    "schedule": schedule
                })
                self.progress.emit(len(out))
        except Exception as e:
            self.error.emit(f"خطأ في جمع المهام بـ schtasks: {e}")
        return out

    def run(self):
        initialized_com = False
        try:
            try:
                pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
            except AttributeError:
                pythoncom.CoInitialize()
            initialized_com = True

            tasks = []
            try:
                if HAVE_COM:
                    tasks = self._collect_tasks_com()
                else:
                    tasks = self._collect_tasks_schtasks()
            except Exception:
                tasks = self._collect_tasks_schtasks()

            total_tasks = len(tasks)
            suspicious_results: List[Dict[str, Any]] = []

            kw_args_lower = [k.strip().lower() for k in split_tokens(self.crit.kw_args)]
            want_sched_lower = self.crit.sched_type.lower()

            for t in tasks:
                if self._stop:
                    break
                reasons = []
                # معيار الاشتباه يعتمد على مطابقة الكلمات المفتاحية فقط
                if kw_args_lower:
                    matched_token = self._match_args_keywords(t.get("command", ""), t.get("arguments", ""),
                                                              kw_args_lower)
                    if matched_token:
                        reasons.append(f"{tr('reason_kw_arg')}: {matched_token}")

                if want_sched_lower != "all" and self._sched_match(t.get("schedule", ""), want_sched_lower):
                    reasons.append(tr("reason_kw_arg"))

                sig_status = check_signature_status(t.get("command", ""))
                if sig_status in ("Invalid", "Unsigned"):
                    reasons.append(tr("reason_sig"))

                is_any_filter_active = bool(kw_args_lower) or (want_sched_lower != "all")
                if is_any_filter_active and reasons:
                    created_val = t.get("created")
                    created_str = created_val.strftime(
                        "%Y-%m-%d %H:%M:%S") if created_val else self._fallback_created_str(t.get("command", ""))
                    exe_name = os.path.basename(t.get("command", "")) or ""
                    args_show = (t.get("arguments", "") or "").strip()
                    if len(args_show) > 80:
                        args_show = args_show[:77] + "..."
                    signature_col = f"{sig_status} | {exe_name} {args_show}".strip()

                    suspicious_results.append({
                        "name": t.get("name", ""),
                        "path": t.get("path", ""),
                        "command": t.get("command", ""),
                        "arguments": t.get("arguments", ""),
                        "author": t.get("author", ""),
                        "created": created_str,
                        "schedule": t.get("schedule", "All"),
                        "signature": signature_col,
                        "reasons": reasons,
                    })

            self.finished.emit(suspicious_results, total_tasks)
        except Exception as e:
            self.error.emit(str(e))
            traceback.print_exc()
        finally:
            if initialized_com:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    def _fallback_created_str(self, exe_path: str) -> str:
        try:
            if exe_path and os.path.exists(exe_path):
                ts = os.path.getctime(exe_path)
                return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            pass
        return tr("unknown")

# ================= بطاقة (Card) =================
class Card(QFrame):
    def __init__(self, title: str):
        super().__init__()
        self.setObjectName("Card")
        self.setFrameShape(QFrame.StyledPanel)
        self.setFrameShadow(QFrame.Raised)
        v_layout = QVBoxLayout(self)
        v_layout.setContentsMargins(16, 16, 16, 16)
        v_layout.setSpacing(10)
        self.title_lbl = QLabel(title)
        self.title_lbl.setObjectName("CardTitle")
        v_layout.addWidget(self.title_lbl)
        self.v = v_layout

# ================= الواجهة الرئيسية =================
class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        self.saved_kw_args: List[str] = []
        self._load_lists()
        self._build_ui()
        self._apply_style_choice()
        self._apply_language()
        self.scanner: Optional[TasksScannerThread] = None
        self.last: List[Dict[str, Any]] = []
        self.showMaximized()  # تشغيل التطبيق بوضع ملء الشاشة بدون تحديد حجم ثابت مسبقًا

    # ---------- بناء الواجهة
    def _build_ui(self):
        self.setWindowTitle(tr("title"))
        # إزالة self.resize لتحديد حجم ثابت، فالشاشة ستملأ تلقائيًا
        root_widget = QWidget()
        self.setCentralWidget(root_widget)
        root_v_layout = QVBoxLayout(root_widget)
        root_v_layout.setContentsMargins(12, 12, 12, 12)
        root_v_layout.setSpacing(10)

        # شريط علوي مع أزرار
        top_buttons_layout = QHBoxLayout()
        top_buttons_layout.setSpacing(10)
        self.btn_scan = QPushButton(tr("scan"))
        self.btn_scan.setObjectName("ActionButton")
        self.btn_stop = QPushButton(tr("stop"))
        self.btn_stop.setObjectName("ActionButton")
        self.btn_stop.setEnabled(False)
        self.btn_clear = QPushButton(tr("clear"))
        self.btn_clear.setObjectName("ActionButton")
        self.btn_exit = QPushButton(tr("exit"))
        self.btn_exit.setObjectName("ActionButton")
        for btn in (self.btn_scan, self.btn_stop, self.btn_clear, self.btn_exit):
            btn.setMinimumHeight(44)
            btn.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Fixed)
        top_buttons_layout.addWidget(self.btn_scan)
        top_buttons_layout.addWidget(self.btn_stop)
        top_buttons_layout.addWidget(self.btn_clear)
        top_buttons_layout.addStretch(1)
        # زر التصدير، تعديل اللون بتعيين نفس اسم الكائن "ActionButton" ليتأثر بالثيم
        self.btn_export = QToolButton()
        self.btn_export.setText(tr("export"))
        self.btn_export.setObjectName("ActionButton")
        # إنشاء قائمة التصدير وربط كل خيار بإجراء محدد مباشرة
        export_menu = QMenu(self)
        act_xlsx = export_menu.addAction(tr("export_xlsx"))
        # act_pdf = export_menu.addAction(tr("export_pdf")) # تم إزالة هذا الخيار
        act_word = export_menu.addAction(tr("export_word"))
        act_txt = export_menu.addAction(tr("export_txt"))
        act_csv = export_menu.addAction(tr("export_csv"))
        act_xlsx.triggered.connect(self._export_all_xlsx)
        # act_pdf.triggered.connect(self._export_pdf) # تم إزالة هذا الربط
        act_word.triggered.connect(self._export_word)
        act_txt.triggered.connect(self._export_txt)
        act_csv.triggered.connect(self._export_selected_csv)
        self.btn_export.setMenu(export_menu)
        self.btn_export.setPopupMode(QToolButton.MenuButtonPopup)
        top_buttons_layout.addWidget(self.btn_export)
        # إزالة زر استعادة خيارات التوقيع حسب المطلوب
        top_buttons_layout.addWidget(self.btn_exit)
        root_v_layout.addLayout(top_buttons_layout)

        # شعار + لغة/ثيم/المستخدم
        logo_lang_theme_layout = QHBoxLayout()
        logo_lang_theme_layout.setSpacing(12)
        self.logo_label = QLabel()
        pm = pixmap_from_base64(SAFE_LOGO_BASE64)
        self.logo_label.setPixmap(pm.scaled(110, 110, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        self.logo_label.setFixedSize(120, 120)
        logo_lang_theme_layout.addWidget(self.logo_label, 0, Qt.AlignLeft | Qt.AlignVCenter)

        selection_box = QFrame()
        selection_box.setObjectName("MiniCard")
        selection_v_layout = QHBoxLayout(selection_box)
        selection_v_layout.setContentsMargins(12, 12, 12, 12)
        selection_v_layout.setSpacing(12)
        self.theme_combo = QComboBox()
        self.theme_combo.addItems([
            "غامق (Dark)", "فاتح (Light)", "أزرق (Blue)", "أخضر (Green)", "أحمر (Red)", "بنفسجي (Purple)"
        ])
        self.user_combo = QComboBox()
        self.user_combo.addItems(["all", "Administrator", "system", "آخر"])
        self.user_line = QLineEdit()
        self.user_line.setPlaceholderText("ادخل اسم المستخدم")
        self.user_line.setVisible(False)
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["العربية", "English"])
        selection_v_layout.addWidget(QLabel(tr("theme")))
        selection_v_layout.addWidget(self.theme_combo)
        selection_v_layout.addSpacing(20)
        selection_v_layout.addWidget(QLabel(tr("lang")))
        selection_v_layout.addWidget(self.lang_combo)
        selection_v_layout.addSpacing(20)
        selection_v_layout.addWidget(QLabel(tr("user")))
        selection_v_layout.addWidget(self.user_combo)
        selection_v_layout.addWidget(self.user_line)
        logo_lang_theme_layout.addWidget(selection_box, 1)
        root_v_layout.addLayout(logo_lang_theme_layout)

        # الشبكة العلوية: تعديل التخطيط ليصبح قسم المدخلات أكبر وقسم الإحصاءات أصغر قليلاً وبشكل متوازي
        upper_grid_layout = QGridLayout()
        upper_grid_layout.setHorizontalSpacing(10)
        upper_grid_layout.setVerticalSpacing(10)
        # تعيين نسب العرض: المدخلات بنسبة 3 والإحصاءات بنسبة 2 (يمكن التعديل حسب الحاجة)
        upper_grid_layout.setColumnStretch(0, 3)
        upper_grid_layout.setColumnStretch(1, 2)
        root_v_layout.addLayout(upper_grid_layout, 2)

        # بطاقة المدخلات
        self.card_inputs = Card(tr("inputs"))
        inputs_v_layout = self.card_inputs.v
        inputs_grid = QGridLayout()
        inputs_grid.setHorizontalSpacing(8)
        inputs_grid.setVerticalSpacing(6)
        inputs_grid.addWidget(QLabel(tr("kw_args")), 0, 0, 1, 4)
        self.kw_filter = QLineEdit()
        self.kw_filter.setPlaceholderText(tr("kw_filter"))
        inputs_grid.addWidget(self.kw_filter, 1, 0, 1, 4)
        self.kw_args = QListWidget()
        for k in self.saved_kw_args:
            self.kw_args.addItem(QListWidgetItem(k))
        self.kw_args.setMinimumHeight(150)
        self.kw_args.setMinimumWidth(480)
        self.kw_args.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        arg_btns_frame = QFrame()
        arg_btns_layout = QVBoxLayout(arg_btns_frame)
        arg_btns_layout.setContentsMargins(0, 0, 0, 0)
        arg_btns_layout.setSpacing(6)
        btn_add_kwarg = QPushButton(tr("add"))
        btn_add_kwarg.setFixedSize(40, 36)
        btn_rem_kwarg = QPushButton(tr("remove"))
        btn_rem_kwarg.setFixedSize(40, 36)
        arg_btns_layout.addWidget(btn_add_kwarg)
        arg_btns_layout.addWidget(btn_rem_kwarg)
        arg_btns_layout.addStretch(1)
        arg_list_h_layout = QHBoxLayout()
        arg_list_h_layout.addWidget(self.kw_args)
        arg_list_h_layout.addWidget(arg_btns_frame)
        inputs_grid.addLayout(arg_list_h_layout, 2, 0, 1, 4)
        inputs_grid.addWidget(QLabel(tr("schedule_type")), 3, 0)
        self.schedule_combo = QComboBox()
        self.schedule_combo.addItems([
            tr("schedule_any"), tr("schedule_daily"), tr("schedule_weekly"), tr("schedule_monthly"), tr("schedule_once")
        ])
        inputs_grid.addWidget(self.schedule_combo, 3, 1)
        inputs_v_layout.addLayout(inputs_grid)
        upper_grid_layout.addWidget(self.card_inputs, 0, 0)

        # بطاقة الإحصاءات: جعل عرضها أصغر قليلاً
        self.card_stats = Card(tr("stats"))
        stats_v_layout = self.card_stats.v
        self.lbl_total = QLabel("0")
        self.lbl_susp = QLabel("0")
        self.lbl_rate = QLabel("0%")
        stats_meta_layout = QHBoxLayout()
        stats_meta_layout.addWidget(QLabel("Total/الإجمالي:"))
        stats_meta_layout.addWidget(self.lbl_total)
        stats_meta_layout.addStretch(1)
        stats_meta_layout.addWidget(QLabel("Suspicious/المشبوه:"))
        stats_meta_layout.addWidget(self.lbl_susp)
        stats_meta_layout.addStretch(1)
        stats_meta_layout.addWidget(QLabel("Rate/النسبة:"))
        stats_meta_layout.addWidget(self.lbl_rate)
        stats_v_layout.addLayout(stats_meta_layout)
        stats_v_layout.addWidget(QLabel(tr("reasons_chart")))
        if pg:
            self.plot = pg.PlotWidget()
            self.plot.setMinimumHeight(160)
            self.plot.setMaximumHeight(200)
            stats_v_layout.addWidget(self.plot)
        else:
            self.plot = None
            lbl = QLabel("(Install pyqtgraph to see chart)")
            stats_v_layout.addWidget(lbl)
        upper_grid_layout.addWidget(self.card_stats, 0, 1)

        # بطاقة النتائج
        self.card_results = Card(tr("results"))
        results_v_layout = self.card_results.v
        self.filter_bar = QLineEdit()
        self.filter_bar.setPlaceholderText(tr("search_bar_placeholder"))
        self.filter_bar.setVisible(False)
        results_v_layout.addWidget(QLabel(tr("filter_hint")))
        results_v_layout.addWidget(self.filter_bar)
        self.table = QTableWidget(0, len(tr("tbl_headers")))
        self.table.setHorizontalHeaderLabels(tr("tbl_headers"))
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.horizontalHeader().setMinimumSectionSize(80)
        self.table.horizontalHeader().setDefaultSectionSize(120)
        self.table.setSortingEnabled(True)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self._show_table_context_menu)
        self.table.cellDoubleClicked.connect(self._show_details_dialog)
        header_height = self.table.horizontalHeader().height()
        self.table.setMinimumHeight(header_height + self.table.verticalHeader().defaultSectionSize() * 12 + 60)
        results_v_layout.addWidget(self.table)
        root_v_layout.addWidget(self.card_results, 5)

        # شريط الحالة
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        self.status.addPermanentWidget(self.progress, 1)

        # ربط الإشارات
        self.btn_scan.clicked.connect(self._start_scan)
        self.btn_stop.clicked.connect(self._stop_scan)
        self.btn_clear.clicked.connect(self._clear)
        # زر التصدير الآن يعتمد على الخيارات داخل قائمة التصدير ولا يحتاج لاتصال آخر
        self.btn_exit.clicked.connect(QApplication.instance().quit)
        self.theme_combo.currentIndexChanged.connect(self._apply_style_choice)
        self.lang_combo.currentIndexChanged.connect(self._on_lang_change)
        self.user_combo.currentIndexChanged.connect(self._on_user_change)
        self.filter_bar.textChanged.connect(self._apply_table_filter)
        self.kw_filter.textChanged.connect(self._filter_kw_args)
        btn_add_kwarg.clicked.connect(self._add_kw_arg)
        btn_rem_kwarg.clicked.connect(self._rem_kw_arg)
        self.shortcut_find = QKeySequence.Find
        self.installEventFilter(self)

    def eventFilter(self, obj, event):
        if event.type() == QEvent.KeyPress and QKeySequence(event.modifiers() | event.key()) == self.shortcut_find:
            self.filter_bar.setVisible(not self.filter_bar.isVisible())
            if self.filter_bar.isVisible():
                self.filter_bar.setFocus()
            return True
        return super().eventFilter(obj, event)

    def _show_message(self, title: str, text_msg: str, is_warning: bool = False):
        try:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle(title)
            msg_box.setText(text_msg)
            msg_box.setIcon(QMessageBox.Warning if is_warning else QMessageBox.Information)
            # تحديث الأنماط لعلب الرسائل لتتناسب مع الثيمات
            if self.theme_combo.currentIndex() == 0: # Dark
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #12183a; } QPushButton{ background: #3152ff; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif self.theme_combo.currentIndex() == 1: # Light
                msg_box.setStyleSheet("QLabel{ color: #003366; } QMessageBox{ background: #ffffff; } QPushButton{ background: #4c6cff; color:#ffffff; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif self.theme_combo.currentIndex() == 2: # Blue
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #1e3d59; } QPushButton{ background: #277da1; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif self.theme_combo.currentIndex() == 3: # Green
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #2e4600; } QPushButton{ background: #486b00; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif self.theme_combo.currentIndex() == 4: # Red
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #67000d; } QPushButton{ background: #8e000f; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif self.theme_combo.currentIndex() == 5: # Purple
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #3d0c62; } QPushButton{ background: #6a0dad; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            else: # Fallback to default dark
                msg_box.setStyleSheet("QLabel{ color: #ffcc00; } QMessageBox{ background: #12183a; } QPushButton{ background: #3152ff; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")

            msg_box.exec_()
        except Exception:
            (QMessageBox.warning if is_warning else QMessageBox.information)(self, title, text_msg)

    def _on_lang_change(self, idx: int):
        global LANG
        LANG = "ar" if idx == 0 else "en"
        self._apply_language()

    def _on_user_change(self, idx: int):
        if self.user_combo.currentText() == "آخر":
            self.user_line.setVisible(True)
        else:
            self.user_line.setVisible(False)

    def _apply_language(self):
        self.setWindowTitle(tr("title"))
        self.btn_scan.setText(tr("scan"))
        self.btn_stop.setText(tr("stop"))
        self.btn_clear.setText(tr("clear"))
        self.btn_export.setText(tr("export"))
        self.btn_exit.setText(tr("exit"))
        self.card_inputs.title_lbl.setText(tr("inputs"))
        self.card_stats.title_lbl.setText(tr("stats"))
        self.card_results.title_lbl.setText(tr("results"))
        self.filter_bar.setPlaceholderText(tr("search_bar_placeholder"))
        self.table.setHorizontalHeaderLabels(tr("tbl_headers"))
        self.schedule_combo.blockSignals(True)
        cur = self.schedule_combo.currentIndex()
        self.schedule_combo.clear()
        self.schedule_combo.addItems([
            tr("schedule_any"), tr("schedule_daily"), tr("schedule_weekly"), tr("schedule_monthly"), tr("schedule_once")
        ])
        self.schedule_combo.setCurrentIndex(cur)
        self.schedule_combo.blockSignals(False)

    def _apply_style_choice(self):
        idx = self.theme_combo.currentIndex()
        if idx == 0:
            self._apply_dark()
        elif idx == 1:
            self._apply_light()
        elif idx == 2:
            self._apply_blue()
        elif idx == 3:
            self._apply_green()
        elif idx == 4:
            self._apply_red()
        elif idx == 5:
            self._apply_purple()
        else:
            self._apply_dark()

    def _apply_dark(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #eaeef8; }
            QMainWindow { background: #0b1020; }
            #Card { background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #141a33, stop:1 #0c1228); border: 1px solid #2a3258; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#ffcc00; }
            #MiniCard { background: qradialgradient(cx:0.2,cy:0.3,radius:1.2, stop:0 #152043, stop:1 #0b1020); border:1px solid #263159; border-radius: 14px; }
            QLabel { color:#ffcc00; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#0b132a; border:1px solid #364272; border-radius:10px; padding:8px; color:#ffcc00; }
            QTableWidget { background:#0b132a; gridline-color:#2a355f; border:1px solid #2a3258; border-radius:14px; }
            QHeaderView::section { background:#18224a; color:#ffcc00; padding:9px; }
            QTableWidget::item:selected { background:#243569; }
            QStatusBar { background:#12183a; border-top:1px solid #2a3258; }
            QProgressBar { background:#0b132a; border:1px solid #364272; border-radius:10px; text-align:center; }
            QProgressBar::chunk { background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #41b3ff, stop:1 #7a5cff); border-radius:10px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #3152ff, stop:1 #b03bff);
                color:#ffcc00;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#1c253f; }
            """
        )
        if self.plot:
            self.plot.setBackground('#0b1020')
            self.plot.getAxis('left').setPen('#ffcc00')
            self.plot.getAxis('bottom').setPen('#ffcc00')
            self.plot.getAxis('left').setTextPen('#ffcc00')
            self.plot.getAxis('bottom').setTextPen('#ffcc00')

    def _apply_light(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #1b233a; }
            QMainWindow { background: #f6f8fe; }
            #Card { background: #ffffff; border: 1px solid #e6ecff; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#003366; }
            #MiniCard { background:#ffffff; border:1px solid #e6ecff; border-radius: 14px; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#ffffff; border:1px solid #cfd8ff; border-radius:10px; padding:8px; color:#003366; }
            QTableWidget { background:#ffffff; gridline-color:#e4eaff; border:1px solid #e6ecff; border-radius:14px; }
            QHeaderView::section { background:#eef3ff; color:#003366; padding:9px; }
            QTableWidget::item:selected { background:#dfe6ff; }
            QStatusBar { background:#eef3ff; border-top:1px solid #dde5ff; }
            QProgressBar { background:#ffffff; border:1px solid #cfd8ff; border-radius:10px; text-align:center; }
            QProgressBar::chunk { background:#4c6cff; border-radius:10px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background:#4c6cff;
                color:#ffffff;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#99ccff; }
            """
        )
        if self.plot:
            self.plot.setBackground('w')
            self.plot.getAxis('left').setPen('#003366')
            self.plot.getAxis('bottom').setPen('#003366')
            self.plot.getAxis('left').setTextPen('#003366')
            self.plot.getAxis('bottom').setTextPen('#003366')

    def _apply_blue(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #ffffff; }
            QMainWindow { background: #1e3d59; }
            #Card { background: #1e3d59; border: 1px solid #277da1; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#ffcc00; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#f4f4f9; border:1px solid #277da1; border-radius:10px; padding:8px; color:#1e3d59; }
            QTableWidget { background:#f4f4f9; gridline-color:#277da1; border:1px solid #277da1; border-radius:14px; }
            QHeaderView::section { background:#277da1; color:#ffffff; padding:9px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background:#277da1;
                color:#ffcc00;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#4d908e; }
            """
        )
        if self.plot:
            self.plot.setBackground('#1e3d59')
            self.plot.getAxis('left').setPen('#ffcc00')
            self.plot.getAxis('bottom').setPen('#ffcc00')
            self.plot.getAxis('left').setTextPen('#ffcc00')
            self.plot.getAxis('bottom').setTextPen('#ffcc00')

    def _apply_green(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #ffffff; }
            QMainWindow { background: #2e4600; }
            #Card { background: #2e4600; border: 1px solid #486b00; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#ffcc00; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#f1f1f1; border:1px solid #486b00; border-radius:10px; padding:8px; color:#2e4600; }
            QTableWidget { background:#f1f1f1; gridline-color:#486b00; border:1px solid #486b00; border-radius:14px; }
            QHeaderView::section { background:#486b00; color:#ffffff; padding:9px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background:#486b00;
                color:#ffcc00;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#99c24d; }
            """
        )
        if self.plot:
            self.plot.setBackground('#2e4600')
            self.plot.getAxis('left').setPen('#ffcc00')
            self.plot.getAxis('bottom').setPen('#ffcc00')
            self.plot.getAxis('left').setTextPen('#ffcc00')
            self.plot.getAxis('bottom').setTextPen('#ffcc00')

    def _apply_red(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #ffffff; }
            QMainWindow { background: #67000d; }
            #Card { background: #67000d; border: 1px solid #8e000f; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#ffcc00; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#fefefe; border:1px solid #8e000f; border-radius:10px; padding:8px; color:#67000d; }
            QTableWidget { background:#fefefe; gridline-color:#8e000f; border:1px solid #8e000f; border-radius:14px; }
            QHeaderView::section { background:#8e000f; color:#ffffff; padding:9px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background:#8e000f;
                color:#ffcc00;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#c72c41; }
            """
        )
        if self.plot:
            self.plot.setBackground('#67000d')
            self.plot.getAxis('left').setPen('#ffcc00')
            self.plot.getAxis('bottom').setPen('#ffcc00')
            self.plot.getAxis('left').setTextPen('#ffcc00')
            self.plot.getAxis('bottom').setTextPen('#ffcc00')

    def _apply_purple(self):
        self.setStyleSheet(
            """
            QWidget { font-family: "Segoe UI", Tahoma; font-size: 11pt; color: #ffffff; }
            QMainWindow { background: #3d0c62; }
            #Card { background: #3d0c62; border: 1px solid #6a0dad; border-radius: 16px; }
            #CardTitle { font-weight:700; font-size: 13pt; color:#ffcc00; }
            QLineEdit, QComboBox, QListWidget, QSpinBox { background:#fefefe; border:1px solid #6a0dad; border-radius:10px; padding:8px; color:#3d0c62; }
            QTableWidget { background:#fefefe; gridline-color:#6a0dad; border:1px solid #6a0dad; border-radius:14px; }
            QHeaderView::section { background:#6a0dad; color:#ffffff; padding:9px; }
            QPushButton#ActionButton, QToolButton#ActionButton {
                background:#6a0dad;
                color:#ffcc00;
                border:0;
                border-radius: 12px;
                padding:12px 18px;
                font-weight: bold; /* تحسين وضوح الخط */
                font-size: 12pt; /* زيادة حجم الخط قليلاً */
            }
            QPushButton#ActionButton:hover, QToolButton#ActionButton:hover { background:#b19cd9; }
            """
        )
        if self.plot:
            self.plot.setBackground('#3d0c62')
            self.plot.getAxis('left').setPen('#ffcc00')
            self.plot.getAxis('bottom').setPen('#ffcc00')
            self.plot.getAxis('left').setTextPen('#ffcc00')
            self.plot.getAxis('bottom').setTextPen('#ffcc00')

    def _add_kw_arg(self):
        text, ok = QInputDialog.getText(self, tr("kw_args"), tr("add"))
        if ok and text.strip():
            current_vals = [self.kw_args.item(i).text() for i in range(self.kw_args.count())]
            if text.strip() not in current_vals:
                self.kw_args.addItem(QListWidgetItem(text.strip()))
            self._save_lists_silent()

    def _rem_kw_arg(self):
        row = self.kw_args.currentRow()
        if row >= 0:
            self.kw_args.takeItem(row)
            self._save_lists_silent()

    def _save_lists_silent(self):
        data = {"kw_args": [self.kw_args.item(i).text() for i in range(self.kw_args.count())]}
        try:
            with open(LISTS_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _load_lists(self):
        if LISTS_FILE.exists():
            try:
                with open(LISTS_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.saved_kw_args = data.get("kw_args", ["ssh", "root", "-L", "@root"])
            except Exception:
                self.saved_kw_args = ["ssh", "root", "-L", "@root"]
        else:
            self.saved_kw_args = ["ssh", "root", "-L", "@root"]

    def _criteria(self) -> Criteria:
        schedule_index = self.schedule_combo.currentIndex()
        schedule_map = {0: "all", 1: "Daily", 2: "Weekly", 3: "Monthly", 4: "Once"}
        selected_schedule_type = schedule_map.get(schedule_index, "all")
        user_val = self.user_combo.currentText()
        if user_val == "آخر":
            user_val = self.user_line.text().strip()
        return Criteria(
            kw_args=[self.kw_args.item(i).text() for i in range(self.kw_args.count())],
            sched_type=selected_schedule_type,
            user=user_val
        )

    def _start_scan(self):
        criteria = self._criteria()
        is_any_filter_active = bool(criteria.kw_args) or (criteria.sched_type.lower() != "all")
        if not is_any_filter_active:
            self._show_message(tr("title"), tr("no_filters"), is_warning=True)
            return
        self.table.setRowCount(0)
        self.last = []
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        self.status.showMessage(tr("progress"))
        self.btn_scan.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.scanner = TasksScannerThread(criteria)
        self.scanner.progress.connect(self._on_progress)
        self.scanner.finished.connect(self._on_finished)
        self.scanner.error.connect(self._on_error)
        self.scanner.start()

    def _stop_scan(self):
        if self.scanner:
            self.scanner.stop()
        self.btn_stop.setEnabled(False)

    def _on_progress(self, n: int):
        pass

    def _on_finished(self, results: List[Dict[str, Any]], total: int):
        self.progress.setVisible(False)
        self.progress.setRange(0, 100)
        self.btn_scan.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.status.showMessage(tr("done").format(len(results), total))
        self.last = results
        self._populate_table(results)
        self._update_stats()

    def _on_error(self, msg: str):
        self.progress.setVisible(False)
        self.btn_scan.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self._show_message("Error", msg, is_warning=True)

    def _clear(self):
        self.table.setRowCount(0)
        self.last = []
        self.lbl_total.setText("0")
        self.lbl_susp.setText("0")
        self.lbl_rate.setText("0%")
        if self.plot:
            self.plot.clear()

    def _populate_table(self, rows: List[Dict[str, Any]]):
        self.table.setSortingEnabled(False)
        headers = tr("tbl_headers")
        self.table.setColumnCount(len(headers))
        self.table.setHorizontalHeaderLabels(headers)
        self.table.setRowCount(len(rows))
        for r, item in enumerate(rows):
            # تعديل أسباب الاشتباه في حال كانت تتضمن سبب مطابقة كلمة (العناصر (Arguments))
            reasons_list = []
            for reason in item.get("reasons", []):
                base = reason.split(":")[0].strip()
                # إذا كان السبب يتطابق مع النص الأصلي، قم بتبديله بالنص الجديد مع إظهار باقي التفاصيل
                if base == tr("reason_kw_arg"):
                    parts = reason.split(":", 1)
                    if len(parts) > 1:
                        reasons_list.append(f"{tr('reason_kw_arg_result')}: {parts[1].strip()}")
                    else:
                        reasons_list.append(tr("reason_kw_arg_result"))
                else:
                    reasons_list.append(reason)
            vals = [
                item.get("name", ""),
                item.get("path", ""),
                item.get("command", ""),
                item.get("arguments", ""),
                item.get("author", ""),
                item.get("created", tr("unknown")),
                item.get("schedule", "All"),
                item.get("signature", "N/A"),
                ", ".join(reasons_list)
            ]
            for c, v in enumerate(vals):
                twi = QTableWidgetItem(str(v))
                twi.setFlags(twi.flags() ^ Qt.ItemIsEditable)
                self.table.setItem(r, c, twi)
        self.table.setSortingEnabled(True)

    def _update_stats(self):
        total = len(self.last)
        susp = 0
        # تحديث الإحصاءات وفقاً للأسباب الموجودة: نحصي سبب مطابقة العناصر (Arguments) وسبب التوقيع
        count_kw = 0
        count_sig = 0
        for it in self.last:
            for r in it.get("reasons", []):
                base = r.split(":")[0].strip()
                if base == tr("reason_kw_arg") or base == tr("reason_kw_arg_result"):
                    count_kw += 1
                if base == tr("reason_sig"):
                    count_sig += 1
            if it.get("reasons"):
                susp += 1
        self.lbl_total.setText(str(total))
        self.lbl_susp.setText(str(susp))
        rate = (susp * 100.0 / total) if total else 0.0
        self.lbl_rate.setText(f"{rate:.1f}%")
        if self.plot:
            self.plot.clear()
            # استخدام المفاتيح الجديدة للإحصاءات
            bars = {tr("reason_kw_arg"): count_kw, tr("reason_sig"): count_sig}
            x = list(range(len(bars)))
            heights = list(bars.values())
            labels = list(bars.keys())
            bg = pg.BarGraphItem(x=x, height=heights, width=0.6, brush='c')
            self.plot.addItem(bg)
            ax = self.plot.getPlotItem().getAxis('bottom')
            ax.setTicks([list(zip(x, labels))])

    def _apply_table_filter(self, text: str):
        text = (text or "").strip().lower()
        for r in range(self.table.rowCount()):
            show = False
            for c in range(self.table.columnCount()):
                it = self.table.item(r, c)
                if it and text in it.text().lower():
                    show = True
                    break
            self.table.setRowHidden(r, not show)

    def _filter_kw_args(self, text: str):
        text = text.strip().lower()
        for i in range(self.kw_args.count()):
            item = self.kw_args.item(i)
            item.setHidden(text not in item.text().lower())

    def _show_table_context_menu(self, pos: QPoint):
        row = self.table.currentRow()
        if row < 0:
            return
        m = QMenu(self)
        act_show_details = m.addAction(tr("show_details")) # إضافة خيار عرض التفاصيل
        act_export_csv = m.addAction(tr("export_csv"))
        act = m.exec_(self.table.viewport().mapToGlobal(pos))
        if act == act_show_details: # ربط خيار عرض التفاصيل بالدالة الموجودة
            self._show_details_dialog(row, -1) # -1 لأن العمود لا يهم هنا
        elif act == act_export_csv:
            self._export_selected_csv()

    def _export_selected_csv(self):
        row = self.table.currentRow()
        if row < 0:
            return
        headers = tr("tbl_headers")
        vals = [self.table.item(row, c).text() if self.table.item(row, c) else "" for c in
                range(self.table.columnCount())]
        name_sanitized = re.sub(r"[^\w\-]+", "_", vals[0] or "selected")
        out_path = APP_DIR / f"{name_sanitized}_task.csv"
        try:
            with open(out_path, "w", encoding="utf-8", newline='') as f:
                f.write(",".join(['"' + h.replace('"', '""') + '"' for h in headers]) + "\n")
                f.write(",".join(['"' + (v or '').replace('"', '""') + '"' for v in vals]) + "\n")
            self._show_message(tr("title"), str(out_path))
        except Exception as e:
            self._show_message(tr("title"), str(e), is_warning=True)

    def _export_all_xlsx(self):
        if not openpyxl:
            self._show_message(tr("title"), tr("need_openpyxl"), is_warning=True)
            return
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Results"
            headers = tr("tbl_headers")
            ws.append(headers)
            for r in range(self.table.rowCount()):
                row_vals = [self.table.item(r, c).text() if self.table.item(r, c) else "" for c in
                            range(self.table.columnCount())]
                ws.append(row_vals)
            for col_idx, _ in enumerate(headers, start=1):
                max_len = 10
                for row in ws.iter_rows(min_col=col_idx, max_col=col_idx):
                    v = row[0].value
                    if v:
                        max_len = max(max_len, len(str(v)))
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 60)
            out_path = APP_DIR / "ntst_results.xlsx"
            wb.save(str(out_path))
            self._show_message(tr("title"), str(out_path))
        except Exception as e:
            self._show_message(tr("title"), str(e), is_warning=True)

    # تم إزالة دالة _export_pdf بالكامل

    def _export_word(self):
        if not Document:
            self._show_message(tr("title"), tr("need_docx"), is_warning=True)
            return
        try:
            out_path = APP_DIR / "ntst_results.docx"
            doc = Document()
            headers = tr("tbl_headers")
            doc.add_paragraph(" | ".join(headers))
            for r in range(self.table.rowCount()):
                row_vals = [self.table.item(r, c).text() if self.table.item(r, c) else "" for c in
                            range(self.table.columnCount())]
                doc.add_paragraph(" | ".join(row_vals))
            doc.save(str(out_path))
            self._show_message(tr("title"), str(out_path))
        except Exception as e:
            self._show_message(tr("title"), str(e), is_warning=True)

    def _export_txt(self):
        try:
            out_path = APP_DIR / "ntst_results.txt"
            with open(out_path, "w", encoding="utf-8") as f:
                headers = tr("tbl_headers")
                f.write(" | ".join(headers) + "\n")
                for r in range(self.table.rowCount()):
                    row_vals = [self.table.item(r, c).text() if self.table.item(r, c) else "" for c in
                                range(self.table.columnCount())]
                    f.write(" | ".join(row_vals) + "\n")
            self._show_message(tr("title"), str(out_path))
        except Exception as e:
            self._show_message(tr("title"), str(e), is_warning=True)

    def _show_details_dialog(self, row: int, col: int): # col parameter is kept for consistency with cellDoubleClicked
        try:
            # التأكد من أن الصف المحدد صالح
            if row < 0 or row >= self.table.rowCount():
                return

            name = self.table.item(row, 0).text() if self.table.item(row, 0) else ""
            path = self.table.item(row, 1).text() if self.table.item(row, 1) else ""
            cmd = self.table.item(row, 2).text() if self.table.item(row, 2) else ""
            args = self.table.item(row, 3).text() if self.table.item(row, 3) else ""
            author = self.table.item(row, 4).text() if self.table.item(row, 4) else ""
            created = self.table.item(row, 5).text() if self.table.item(row, 5) else tr("unknown")
            sched = self.table.item(row, 6).text() if self.table.item(row, 6) else ""
            sig = self.table.item(row, 7).text() if self.table.item(row, 7) else ""
            reasons = self.table.item(row, 8).text() if self.table.item(row, 8) else ""
            if created == tr("unknown") and cmd:
                created = TasksScannerThread(Criteria())._fallback_created_str(cmd)
            text = (
                f"{tr('tbl_headers')[0]}: {name}\n"
                f"{tr('tbl_headers')[1]}: {path}\n"
                f"{tr('tbl_headers')[2]}: {cmd}\n"
                f"{tr('tbl_headers')[3]}: {args}\n"
                f"{tr('tbl_headers')[4]}: {author}\n"
                f"{tr('created')}: {created}\n"
                f"{tr('tbl_headers')[6]}: {sched}\n"
                f"{tr('tbl_headers')[7]}: {sig}\n"
                f"{tr('tbl_headers')[8]}: {reasons}\n"
            )
            dlg = QDialog(self)
            dlg.setWindowTitle(tr("details"))
            # تطبيق الأنماط على مربع الحوار ليتناسب مع الثيم الحالي
            current_theme_idx = self.theme_combo.currentIndex()
            if current_theme_idx == 0: # Dark
                dlg.setStyleSheet("QDialog { background-color: #0b1020; color: #eaeef8; } QLabel { color: #ffcc00; } QTextEdit { background-color: #0b132a; color: #ffcc00; border: 1px solid #364272; border-radius: 8px; } QPushButton { background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #3152ff, stop:1 #b03bff); color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif current_theme_idx == 1: # Light
                dlg.setStyleSheet("QDialog { background-color: #f6f8fe; color: #1b233a; } QLabel { color: #003366; } QTextEdit { background-color: #ffffff; color: #1b233a; border: 1px solid #cfd8ff; border-radius: 8px; } QPushButton { background:#4c6cff; color:#ffffff; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif current_theme_idx == 2: # Blue
                dlg.setStyleSheet("QDialog { background-color: #1e3d59; color: #ffffff; } QLabel { color: #ffcc00; } QTextEdit { background-color: #f4f4f9; color: #1e3d59; border: 1px solid #277da1; border-radius: 8px; } QPushButton { background:#277da1; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif current_theme_idx == 3: # Green
                dlg.setStyleSheet("QDialog { background-color: #2e4600; color: #ffffff; } QLabel { color: #ffcc00; } QTextEdit { background-color: #f1f1f1; color: #2e4600; border: 1px solid #486b00; border-radius: 8px; } QPushButton { background:#486b00; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif current_theme_idx == 4: # Red
                dlg.setStyleSheet("QDialog { background-color: #67000d; color: #ffffff; } QLabel { color: #ffcc00; } QTextEdit { background-color: #fefefe; color: #67000d; border: 1px solid #8e000f; border-radius: 8px; } QPushButton { background:#8e000f; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            elif current_theme_idx == 5: # Purple
                dlg.setStyleSheet("QDialog { background-color: #3d0c62; color: #ffffff; } QLabel { color: #ffcc00; } QTextEdit { background-color: #fefefe; color: #3d0c62; border: 1px solid #6a0dad; border-radius: 8px; } QPushButton { background:#6a0dad; color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")
            else: # Fallback to default dark
                dlg.setStyleSheet("QDialog { background-color: #0b1020; color: #eaeef8; } QLabel { color: #ffcc00; } QTextEdit { background-color: #0b132a; color: #ffcc00; border: 1px solid #364272; border-radius: 8px; } QPushButton { background: qlineargradient(x1:0,y1:0,x2:1,y2:1, stop:0 #3152ff, stop:1 #b03bff); color:#ffcc00; border:0; border-radius: 8px; padding: 6px 12px; }")

            v = QVBoxLayout(dlg)
            te = QTextEdit()
            te.setReadOnly(True)
            te.setPlainText(text)
            v.addWidget(te)
            okb = QPushButton(tr("ok"))
            okb.clicked.connect(dlg.accept)
            h = QHBoxLayout()
            h.addStretch(1)
            h.addWidget(okb)
            v.addLayout(h)
            dlg.resize(720, 480)
            dlg.exec_()
        except Exception as e:
            self._show_message(tr("title"), str(e), is_warning=True)

    # تم إزالة دالة استعادة خيارات التوقيع وفق المتطلبات

# ================== تشغيل التطبيق ==================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    w = Main()
    w.show()
    sys.exit(app.exec_())